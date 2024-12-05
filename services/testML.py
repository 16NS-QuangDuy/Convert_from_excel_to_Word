# encoding: utf-8
import yaml
import os
import copy
import sys
import importlib
import shutil
import traceback
from shutil import copytree, ignore_patterns, move, rmtree
import filecmp
from filecmp import dircmp
from shutil import copyfile
from settings.config import Config
from services.reporter import Reporter
from services.folder import Folder
from services.excelML import ExcelML
from services.capture import Capture
from services.wordML import WordML
from services.elmHelper import ElmHelper
from docx import Document
import win32com.client as win32
import time


class TestResult:

    config_file = Config.get_default_config_file(__file__)

    def __init__(self, unit_test_class_name, test_db_file):
        self.auto_worker_name = self.__class__.__name__
        self.TESTList = []
        self.PASS = 0
        self.FAIL = 0
        self.FAIL_LIST = []
        self.PASS_LIST = []
        self.ERROR_LIST = []
        self.WARN_LIST = []
        self.INFO_LIST = []
        self.config = Config()
        self.config.UnitTestClassName = os.path.basename(unit_test_class_name).replace(".py", "")
        self.config.TestDBYamlFile = test_db_file
        self.config.the_root_dict = dict()
        self.config.TestClassObject = None
        self.config.TestModule = None
        self.config.MyClass = None
        self.config.ClassName = self.config.UnitTestClassName.replace("Test", "").replace("Methods", "")
        self.config.TestID = ""
        self.config.METHOD = ""
        self.config.INIT = ""
        self.config.INPUT = ""
        self.config.OUTPUT = ""
        self.config.RESULT = dict()
        self.config.RETURN = None
        self.config.DebugMethod = None
        self.config.DebugTCID = None
        self.capture = Capture()
        self.capture.start()
        self.capture.stop()
        Config.set_attr_from_yaml(self.__dict__["config"], self.config_file, self.auto_worker_name)

    def execute_auto_task(self):
        self.read_config_info()
        self.load_test_module()
        self.get_tc_list()
        self.execute_customized_auto_task()
        self.command_line_report_test_result()

    def read_config_info(self):
        config_file = self.config.TestDBYamlFile
        with open(config_file, 'r', encoding="utf-8") as stream:
            the_root_dict = yaml.safe_load(stream)
            self.config.the_root_dict = the_root_dict[self.config.UnitTestClassName]

    def load_test_module(self):
        sys.path.insert(0, '')
        self.config.TestModule = importlib.import_module(self.config.the_root_dict["test_module_name"])
        self.config.MyClass = getattr(self.config.TestModule, self.config.the_root_dict["test_class_name"])

    def config_server_base_dir(self, srv_base_dir=None):
        self.config.set_server_base_dir(srv_base_dir)

    def execute_customized_auto_task(self):
        """
        This function is to test all the function list defined in the yaml test file
        :return:
        """
        count = 0
        for test_func_name in getattr(self.config, "the_root_dict")["test_function_list"]:
            self.config.METHOD = test_func_name.replace("test_", "")
            if self.config.DebugMethod is not None and self.config.METHOD != self.config.DebugMethod:
                continue
            the_test_data_dict = getattr(self.config, "the_root_dict")[test_func_name]
            for tc_id, tc_data_dict in the_test_data_dict.items():
                self.config.TestID = tc_id
                self.config.TestDataDict = tc_data_dict
                self.config.INIT = TestResult.get_tc_init_data(tc_data_dict, "init")
                if self.config.DebugTCID is not None and self.config.TestID not in self.config.DebugTCID:
                    continue
                count += 1
                print("         executing testcase...(%s/%s)" % (count, len(self.TESTList)), test_func_name, tc_id)
                if len(tc_data_dict["input"]) == 0:
                    self.config.INPUT = []
                    self.config.OUTPUT = TestResult.get_tc_data_i(tc_data_dict, "output", 0)
                    self.config.RETURN = None
                    try:
                        self.execute_auto_task_one_input()
                    except Exception as e:
                        self.capture.stop()
                        print("        WARNING: " + self.config.TestID + str(e))
                        traceback.print_exc(file=sys.stdout)
                        self.config.ACTUAL = str(e)
                        pass
                    self.create_tc_result()
                    self.add_test_result()
                else:
                    for i, input_i in enumerate(tc_data_dict[tc_data_dict["input"][0]]):
                        self.config.INPUT = TestResult.get_tc_data_i(tc_data_dict, "input", i)
                        self.config.OUTPUT = TestResult.get_tc_data_i(tc_data_dict, "output", i)
                        self.config.RETURN = None
                        try:
                            self.execute_auto_task_one_input()
                        except Exception as e:
                            self.capture.stop()
                            print ("        WARNING: " + self.config.TestID + str(e))
                            traceback.print_exc(file=sys.stdout)
                            self.config.ACTUAL = str(e)
                            pass
                        self.create_tc_result()
                        self.add_test_result()

    def execute_auto_task_one_input(self):
        self.config.RETURN = None
        self.config.TestClassObject = self.config.MyClass()
        self.capture.start()
        if isinstance(self.config.INPUT, list):
            actual = getattr(self.config.TestClassObject, self.config.METHOD)(*self.config.INPUT)
        else:
            actual = getattr(self.config.TestClassObject, self.config.METHOD)(self.config.INPUT)
        actual = TestResult.get_tc_result(self.config.TestDataDict, actual)
        self.capture.stop()
        self.update_error_warn_info()
        self.config.ACTUAL = actual
        self.config.ACTUAL = actual

    def create_tc_result(self):
        result = dict()
        result['MODULE'] = self.config.the_root_dict["test_module_name"]
        result['CLASS'] = self.config.ClassName
        result['METHOD'] = self.config.METHOD
        result['PARAMETER'] = self.config.TestDataDict["input"]
        result['TestID'] = self.config.TestID
        result['INIT'] = self.config.INIT
        result['INPUT'] = self.config.INPUT
        result['OUTPUT_NAME'] = self.config.TestDataDict["output"]
        result['OUTPUT'] = self.config.OUTPUT
        result['ACTUAL'] = self.config.ACTUAL
        if isinstance(self.config.INPUT, list):
            result['INPUT'] = []
            for li in self.config.INPUT:
                if isinstance(li, dict):
                    for k, v in copy.copy(li).items():
                        if k.startswith("__"): del li[k]
                result['INPUT'].append(li)
        elif isinstance(self.config.INPUT, dict):
            for k, v in copy.copy(self.config.INPUT).items():
                if k.startswith("__"): del self.config.INPUT[k]
        self.config.RESULT = result
        return result

    def add_test_result(self):
        if self.config.RETURN is None:
            if self.config.RESULT["OUTPUT"] == self.config.ACTUAL:
                self.config.RETURN = True
            else:
                self.config.RETURN = False
        if self.config.RETURN:
            self.PASS += 1
            self.PASS_LIST.append([self.config.RESULT, self.config.ACTUAL, "PASS"])
        else:
            self.FAIL += 1
            self.FAIL_LIST.append([self.config.RESULT, self.config.ACTUAL, "FAIL"])

    @staticmethod
    def get_tc_data_i(tc_data_dict, tc_data, i):
        if tc_data not in tc_data_dict:
            return None
        tc_data_values = []
        for data_type in tc_data_dict[tc_data]:
            data_i = tc_data_dict[data_type][i]
            if isinstance(data_i, str) and data_i in tc_data_dict:
                data_i = tc_data_dict[data_i]
            tc_data_values.append(data_i)
        if len(tc_data_dict[tc_data]) == 1:
            tc_data_values = tc_data_values[0]
        return tc_data_values

    @staticmethod
    def get_tc_init_data(tc_data_dict, tc_data):
        if tc_data not in tc_data_dict:
            return None
        else:
            return tc_data_dict[tc_data]

    @staticmethod
    def get_tc_result(tc_data_dict, result):
        out_values = []
        for out_type in tc_data_dict["output"]:
            out_value = result
            if isinstance(result, dict):
                if out_type in result:
                    out_value = result[out_type]
            if isinstance(result, list):
                if len(result) > 0 and isinstance(result[0], dict):
                    out_value = [rc_dict[out_type] for rc_dict in result]
            out_values.append(out_value)
        if len(tc_data_dict["output"]) <= 1:
            out_values = out_values[0]
        return out_values

    def command_line_report_test_result(self):
        Reporter.create_command_window_report(self.__class__.__name__, self.PASS_LIST, "test case", "PASS")
        Reporter.create_command_window_report(self.__class__.__name__, self.FAIL_LIST, "test case", "FAIL")
        Reporter.create_command_window_report(self.__class__.__name__, self.ERROR_LIST, "test case", "ERROR")
        Reporter.create_command_window_report(self.__class__.__name__, self.WARN_LIST, "test case", "WARN")
        Reporter.create_command_window_report(self.__class__.__name__, self.INFO_LIST, "test case", "INFO")

    def get_fail_list(self):
        return self.FAIL_LIST

    def get_pass_list(self):
        return self.PASS_LIST

    def config_dbg_tc(self, method_tc_id):
        method_tc_id = method_tc_id.split(" ")
        self.config.DebugMethod = method_tc_id[0]
        if len(method_tc_id) >= 2:
            self.config.DebugTCID = [method_tc_id[1]]

    def config_dbg_tc_list(self, method_tc_id, tc_list):
        self.config.DebugMethod = method_tc_id
        self.config.DebugTCID = tc_list

    def update_error_warn_info(self):
        text = self.capture.get_text()
        if "ERROR" in text or "error" in text:
            error = "%s.%s %s %s" % (self.config.ClassName, self.config.METHOD, self.config.TestID, text)
            self.ERROR_LIST.append(error)
        if "WARN" in text:
            warn = "%s.%s %s %s" % (self.config.ClassName, self.config.METHOD, self.config.TestID, text)
            self.WARN_LIST.append(warn)
        if "INFO" in text:
            info = "%s.%s %s %s" % (self.config.ClassName, self.config.METHOD, self.config.TestID, text)
            self.INFO_LIST.append(info)

    def get_tc_list(self):
        tc_list = []
        for test_func_name in getattr(self.config, "the_root_dict")["test_function_list"]:
            self.config.METHOD = test_func_name.replace("test_", "")
            if self.config.DebugMethod is not None and self.config.METHOD != self.config.DebugMethod:
                continue
            the_test_data_dict = getattr(self.config, "the_root_dict")[test_func_name]
            for tc_id, tc_data_dict in the_test_data_dict.items():
                if self.config.DebugTCID is not None and tc_id not in self.config.DebugTCID:
                    continue
                tc_list.append({"test_func_name": test_func_name, "tc_id": tc_id})
        self.TESTList = tc_list
        return tc_list

    @staticmethod
    def load_test_module2(worker_py, class_name):
        """
        load module and class from input string
        :param module_name: module path
        :param class_name: class name
        :return: class object
        """
        worker_py = worker_py.replace(".\\", "/").replace("\\", "/")
        worker_py = worker_py.replace(Config.BASE_DIR.replace("\\", "/"), "")
        worker_py = worker_py.lstrip("/")
        worker_py = worker_py.rstrip(".py")
        worker_py = worker_py.replace("/", ".")
        module_object = importlib.import_module(worker_py)
        MyClass = getattr(module_object, class_name)
        return MyClass


class WorkerTestResult(TestResult):

    config_file = Config.get_default_config_file(__file__)

    def execute_auto_task_one_input(self):
        # STEP1 Configure the input output
        input_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_input"], self.config.INIT)
        output_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_output"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        actual_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_actual"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        config = Config()
        config.update_config(input_dir=input_dir)
        config.update_config(output_dir=actual_dir)
        config.update_config(InputDir=input_dir)
        config.refresh_config()
        # STEP2 run the method
        self.capture.start()
        self.config.TestClassObject = self.config.MyClass(config)
        self.config.RETURN = None
        Folder.create_directory(actual_dir, exist_del=True)
        if isinstance(self.config.INPUT, list) and self.config.INPUT == []:
            getattr(self.config.TestClassObject, self.config.METHOD)()
        elif isinstance(self.config.INPUT, str):
            getattr(self.config.TestClassObject, self.config.METHOD)(self.config.INPUT)
        self.update_error_warn_info()
        self.config.ACTUAL = actual_dir
        self.config.OUTPUT = output_dir
        self.capture.stop()

    def add_test_result(self):
        if self.config.RETURN is not None:
            actual = Folder.get_all_files(self.config.ACTUAL, "*.xlsx")
            output = Folder.get_all_files(self.config.OUTPUT, "*.xlsx")
            match = 0
            key_col_list = [self.config.ColumnOrdDict[i] for i in self.config.CompareColumnList]
            diff_a = []
            diff_b = []
            for i in range(min(len(actual), len(output))):
                result_dict = ExcelML.diff_excel_a_b(output[i], actual[i], self.config.SheetName, key_col_list)
                if result_dict["result"]:
                    match += 1
                else:
                    diff_a.append(output[i])
                    diff_a.extend(result_dict["diff_list_a"])
                    diff_a.append(actual[i])
                    diff_b.extend(result_dict["diff_list_b"])
            if match > 0 and match == len(output):
                self.config.RETURN = True
        if self.config.RETURN:
            self.PASS += 1
            self.PASS_LIST.append([self.config.RESULT, self.config.ACTUAL, "PASS"])
        else:
            self.FAIL += 1
            self.FAIL_LIST.append([self.config.RESULT, self.config.ACTUAL, "FAIL"])


class UMImageCmpWorkerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def add_test_result(self):
        if not self.config.RETURN:
            actual = Folder.scan_all_files(self.config.ACTUAL)
            output = Folder.scan_all_files(self.config.OUTPUT)
            actual += Folder.scan_all_files("%s/*/" % self.config.ACTUAL)
            output += Folder.scan_all_files("%s/*/" % self.config.OUTPUT)
            actual += Folder.scan_all_files("%s/*/*/" % self.config.ACTUAL)
            output += Folder.scan_all_files("%s/*/*/" % self.config.OUTPUT)
            actual = [file for file in actual if not file.endswith(".bak")]
            output = [file for file in output if not file.endswith(".bak")]
            match = 0
            match_name = 0
            for i in range(min(len(actual), len(output))):
                filecmp.clear_cache()
                if filecmp.cmp(output[i], actual[i]):
                    match += 1
                if os.path.basename(output[i]) == os.path.basename(actual[i]):
                    match_name += 1
            if match > 0 and (len(actual) == len(output) == match == match_name) and match_name > 0:
                self.config.RETURN = True
        if self.config.RETURN:
            self.PASS += 1
            self.PASS_LIST.append([self.config.RESULT, self.config.ACTUAL, "PASS"])
        else:
            self.FAIL += 1
            self.FAIL_LIST.append([self.config.RESULT, self.config.ACTUAL, "FAIL"])


class ManagerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def execute_auto_task_one_input(self):
        # STEP1 Configure the input output
        input_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_input"], self.config.INIT)
        output_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_output"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        actual_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_actual"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        config = Config()
        config.update_config(InputDir=input_dir)
        config.refresh_config()
        Folder.create_directory(actual_dir, exist_del=True)
        # STEP2 run the method
        self.capture.start()
        self.config.TestClassObject = self.config.MyClass(config)
        self.config.RETURN = None
        if isinstance(self.config.INPUT, list) and self.config.INPUT == []:
            getattr(self.config.TestClassObject, self.config.METHOD)()
        elif isinstance(self.config.INPUT, str):
            getattr(self.config.TestClassObject, self.config.METHOD)(self.config.INPUT)
        # copy data from session to actual result
        config = getattr(self.config.TestClassObject, "get_config")()
        file_lists = Folder.scan_all_files(os.path.join(config.OutputDir,
                                                        config.StepOutPutDir[config.StepList[-1]]))
        for file in file_lists:
            src = file
            dst = os.path.join(actual_dir, os.path.basename(file))
            copyfile(src, dst)
        folder_lists = Folder.scan_all_folder(os.path.join(config.OutputDir,
                                                           config.StepOutPutDir[config.StepList[-1]]))
        for folder in folder_lists:
            src = folder
            dst = os.path.join(actual_dir, os.path.basename(folder))
            copytree(src, dst)
        self.capture.stop()
        self.update_error_warn_info()
        self.config.ACTUAL = actual_dir
        self.config.OUTPUT = output_dir


class ExcelWorkerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def add_test_result(self):
        self.config.RETURN = None
        actual = Folder.get_all_files(self.config.ACTUAL, "*.xlsx")
        output = Folder.get_all_files(self.config.OUTPUT, "*.xlsx")
        if len(actual) == len(output):
            self.config.RETURN = False
            match = 0
            diff_a = []
            diff_b = []
            for i in range(min(len(actual), len(output))):
                result_dict = ExcelML.diff_excel_a_b_all(output[i], actual[i])
                if result_dict["result"]:
                    match += 1
                else:
                    diff_a.append(output[i])
                    diff_a.extend(result_dict["diff_list_a"])
                    diff_a.append(actual[i])
                    diff_b.extend(result_dict["diff_list_b"])
            if match == len(output) > 0 >= len(diff_a):
                self.config.RETURN = True
            else:
                self.config.OUTPUT = "\n".join(diff_a)
                self.config.ACTUAL = "\n".join(diff_b)
                self.config.RESULT["OUTPUT"] = self.config.OUTPUT
        if self.config.RETURN:
            self.PASS += 1
            self.PASS_LIST.append([self.config.RESULT, self.config.ACTUAL, "PASS"])
        else:
            self.FAIL += 1
            self.FAIL_LIST.append([self.config.RESULT, self.config.ACTUAL, "FAIL"])


class WordExcelWorkerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def __init__(self, unit_test_class_name, test_db_file):
        self.auto_worker_name = self.__class__.__name__
        WorkerTestResult.__init__(self, unit_test_class_name, test_db_file)
        self.WordApplication = win32.gencache.EnsureDispatch("Word.Application")
        self.WordActualFile = None
        self.WordExpectFile = None
        self.WordActualObj = None
        self.WordExpectObj = None
        self.target_docx = None
        self.ExcelApplication = win32.gencache.EnsureDispatch("Excel.Application")
        self.ExcelActualFile = None
        self.ExcelExpectFile = None
        self.ExcelActualObj = None
        self.ExcelExpectObj = None
        self.target_excel = None
        self.ActualDir = None
        self.OutputDir = None

    def add_test_result(self):
        if self.config.RETURN is None:
            self.verify_excel_file()
            self.verify_word_file()
            self.create_diff_files()
        if self.config.RETURN:
            self.PASS += 1
            self.PASS_LIST.append([self.config.RESULT, self.config.ACTUAL, "PASS"])
        else:
            self.FAIL += 1
            self.FAIL_LIST.append([self.config.RESULT, self.config.ACTUAL, "FAIL"])

    def create_diff_files(self):
        if not self.config.RETURN:
            try:
                self.create_excel_diff_file()
                self.create_comparision_file()
            except:
                try:
                    self.WordApplication.Quit()
                    self.ExcelApplication.Quit()
                except:
                    pass
                pass

    def verify_excel_file(self):
        self.ActualDir = self.config.ACTUAL
        self.OutputDir = self.config.OUTPUT
        self.ExcelActualFile = None
        self.ExcelExpectFile = None
        self.target_excel = None
        actual = Folder.get_all_files(self.ActualDir, "*.xlsx")
        output = Folder.get_all_files(self.OutputDir, "*.xlsx")
        if len(actual) == len(output) > 0:
            self.config.RETURN = False
            match = 0
            diff_a = []
            diff_b = []
            for i in range(min(len(actual), len(output))):
                self.ExcelActualFile = actual[i]
                self.ExcelExpectFile = output[i]
                self.target_excel = actual[i][:-5] + "_diff.xlsx"
                result_dict = ExcelML.diff_excel_a_b_all(output[i], actual[i])
                if result_dict["result"]:
                    match += 1
                else:
                    diff_a.append(output[i])
                    diff_a.extend(result_dict["diff_list_a"])
                    diff_a.append(actual[i])
                    diff_b.extend(result_dict["diff_list_b"])
            if match == len(output) > 0 >= len(diff_a):
                self.config.RETURN = True
            else:
                self.config.OUTPUT = "\n".join(diff_a)
                self.config.ACTUAL = "\n".join(diff_b)
                self.config.RESULT["OUTPUT"] = self.config.OUTPUT
        else:
            self.config.RETURN = False
            self.config.RESULT["OUTPUT"] = "Excel file is created"
            self.config.ACTUAL = "Excel file is not exist"
        self.config.ACTUAL = self.ActualDir
        self.config.OUTPUT = self.OutputDir

    def verify_word_file(self):
        self.ActualDir = self.config.ACTUAL
        self.OutputDir = self.config.OUTPUT
        self.WordActualFile = None
        self.WordExpectFile = None
        self.target_docx = None
        actual = Folder.get_all_files(self.ActualDir, "*.docx")
        output = Folder.get_all_files(self.OutputDir, "*.docx")
        if len(actual) == len(output) > 0:
            self.config.ACTUAL = actual[0]
            self.config.OUTPUT = output[0]
            self.WordActualFile = self.config.ACTUAL
            self.WordExpectFile = self.config.OUTPUT
            self.target_docx = self.WordActualFile[:-5] + "_Comparsion.docx"
            output_records = []
            if os.path.exists(self.config.OUTPUT):
                block_list = [block for block in WordML.iter_block_items_external(Document(self.config.OUTPUT))]
                output_records = self.get_block_format_text(block_list)
            actual_records = []
            if os.path.exists(self.config.ACTUAL):
                block_list = [block for block in WordML.iter_block_items_external(Document(self.config.ACTUAL))]
                actual_records = self.get_block_format_text(block_list)
            output_text = "\n".join(output_records)
            actual_text = "\n".join(actual_records)
            if output_text == actual_text:
                self.config.RETURN = True
            else:
                self.config.RETURN = False
                self.config.RESULT["OUTPUT"] = self.config.OUTPUT = "Check comparison file"
                self.config.ACTUAL = "Check comparison file"
        else:
            self.config.RETURN = False
            self.config.RESULT["OUTPUT"] = "Word file is created"
            self.config.ACTUAL = "Word file is not exist"
        self.config.ACTUAL = self.ActualDir
        self.config.OUTPUT = self.OutputDir

    def verify_word_header_footer(self):
        actual = Folder.get_all_files(self.config.ACTUAL, "*.docx")
        output = Folder.get_all_files(self.config.OUTPUT, "*.docx")
        if len(actual) == len(output) > 0:
            self.config.ACTUAL = actual[0]
            self.config.OUTPUT = output[0]
            output_records = []
            if os.path.exists(self.config.OUTPUT):
                doc_output = Document(self.config.OUTPUT)
                section_output = doc_output.sections[0]
                lst_txbx_header_output = ElmHelper.find_textbox(section_output.header._element)['t_lst']
                lst_txbx_footer_output = ElmHelper.find_textbox(section_output.footer._element)['t_lst']
                result_txbx_output = lst_txbx_header_output + lst_txbx_footer_output
                header_output = [header.text for header in section_output.header.paragraphs]
                header_output_even_output = [header.text for header in section_output.even_page_header.paragraphs]
                header_output_first_output = [header.text for header in section_output.first_page_header.paragraphs]
                footer_output = [footer.text for footer in section_output.footer.paragraphs]
                footer_output_even_output = [footer.text for footer in section_output.even_page_footer.paragraphs]
                footer_output_first_output = [footer.text for footer in section_output.first_page_footer.paragraphs]
                output_records = header_output + header_output_even_output + header_output_first_output + footer_output + footer_output_even_output + footer_output_first_output + result_txbx_output
            actual_records = []
            if os.path.exists(self.config.ACTUAL):
                doc_actual = Document(self.config.ACTUAL)
                section_actual = doc_actual.sections[0]
                lst_txbx_header_actual = ElmHelper.find_textbox(section_actual.header._element)['t_lst']
                lst_txbx_footer_actual = ElmHelper.find_textbox(section_actual.footer._element)['t_lst']
                result_txbx_actual = lst_txbx_header_actual + lst_txbx_footer_actual
                header_actual = [header.text for header in section_actual.header.paragraphs]
                header_actual_even_output = [header.text for header in section_actual.even_page_header.paragraphs]
                header_actual_first_output = [header.text for header in section_actual.first_page_header.paragraphs]
                footer_actual = [footer.text for footer in section_actual.footer.paragraphs]
                footer_actual_even_output = [footer.text for footer in section_actual.even_page_footer.paragraphs]
                footer_actual_first_output = [footer.text for footer in section_actual.first_page_footer.paragraphs]
                actual_records = header_actual + header_actual_even_output + header_actual_first_output + footer_actual + footer_actual_even_output + footer_actual_first_output + result_txbx_actual
            output_text = "\n".join(output_records)
            actual_text = "\n".join(actual_records)
            if output_text == actual_text:
                self.config.RETURN = True
            else:
                self.config.RETURN = False
                self.config.RESULT["OUTPUT"] = self.config.OUTPUT = output_text
                self.config.ACTUAL = actual_text
        else:
            self.config.RETURN = False
            self.config.RESULT["OUTPUT"] = "Word file is created"
            self.config.ACTUAL = "Word file is not exist"

    @staticmethod
    def get_block_format_text(block_list):
        records = []
        WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for block in block_list:
            if WordML.is_paragraph_block(block):
                records.append(block.text)
            elif WordML.is_table_block(block):
                table = block
                for tr_i in ElmHelper.findall(table._tbl, 'tr'):
                    row_cells = ElmHelper.get_tr_text(tr_i)
                    records.append(",".join([text for text in row_cells]))
                    all_tcw = tr_i.findall(".//w:tc/w:tcPr/w:tcW", namespaces=tr_i.nsmap)
                    tcw_lst = [str(w.attrib[WORD_NAMESPACE + 'w']) for w in all_tcw]
                    records.append(",".join(tcw_lst))
        return records

    def create_comparision_file(self):
        if self.WordActualFile is None or self.WordExpectFile is None:
            return
        self.WordActualObj = self.WordApplication.Documents.Open(self.WordActualFile)
        self.WordExpectObj = self.WordApplication.Documents.Open(self.WordExpectFile)
        self.WordApplication.CompareDocuments(self.WordExpectObj,
                                              self.WordActualObj,
                                              RevisedAuthor=False,
                                              IgnoreAllComparisonWarnings=True)
        # before saving if you like viewing the document in Print Layout.
        # Otherwise the saved Comparison.docx opens as Web Layout by default (Type = 6).
        self.WordApplication.ActiveDocument.ActiveWindow.View.Type = 3
        # Save the comparison document as "Comparison.docx"
        self.WordApplication.ActiveDocument.SaveAs(FileName=self.target_docx, Encoding=20127)
        self.WordApplication.ActiveDocument.Close()
        self.WordActualObj.Close()
        self.WordExpectObj.Close()
        time.sleep(1)

    def create_excel_diff_file(self):
        if self.ExcelActualFile is None or self.ExcelExpectFile is None:
            return
        path1 = self.ExcelActualFile[:-5] + "_Diff.xlsx"
        path2 = self.ExcelExpectFile
        copyfile(self.ExcelActualFile, path1)
        # os.rename(self.ExcelActualFile, path1)
        self.ExcelActualObj = self.ExcelApplication.Workbooks.Open(Filename=path1)
        self.ExcelExpectObj = self.ExcelApplication.Workbooks.Open(Filename=path2)
        ws2 = self.ExcelExpectObj.Worksheets(1)
        ws2.Copy(After=self.ExcelActualObj.Worksheets(1))
        self.ExcelActualObj.Worksheets(1).Name = "Actual"
        self.ExcelActualObj.Worksheets(2).Name = "Output"
        self.ExcelActualObj.Close(SaveChanges=True)
        self.ExcelExpectObj.Close()


class NightlyTestManagerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def execute_auto_task_one_input(self):
        # STEP1 Configure the input output
        self.config.RETURN is None
        input_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_input"], self.config.INIT)
        output_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_output"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        actual_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_actual"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        config = Config()
        tc_output_dir = os.path.join(self.config.SRV_BASE_DIR, "test", "Output", self.config.ClassName, self.config.TestID)
        config.update_config(InputDir=input_dir)
        config.update_config(OutputDir=tc_output_dir)
        config.update_config(input_dir=input_dir)
        config.refresh_config()
        Folder.create_directory(actual_dir, exist_del=True)
        # STEP2 run the method
        self.capture.start()
        self.config.TestClassObject = self.config.MyClass(config)
        copied_config = getattr(self.config.TestClassObject, "get_config")()
        self.config.RETURN = None
        if isinstance(self.config.INPUT, list) and self.config.INPUT == []:
            getattr(self.config.TestClassObject, self.config.METHOD)()
        elif isinstance(self.config.INPUT, str):
            getattr(self.config.TestClassObject, self.config.METHOD)(self.config.INPUT)
        # copy data from session to actual result
        file_lists = Folder.scan_all_files(copied_config.StepOutPutDir[copied_config.StepList[-1]])
        for file in file_lists:
            src = file
            dst = os.path.join(actual_dir, os.path.basename(file))
            copyfile(src, dst)
        folder_lists = Folder.scan_all_folder(copied_config.StepOutPutDir[copied_config.StepList[-1]])
        for folder in folder_lists:
            src = folder
            dst = os.path.join(actual_dir, os.path.basename(folder))
            copytree(src, dst)
        self.capture.stop()
        self.update_error_warn_info()
        self.config.ACTUAL = actual_dir
        self.config.OUTPUT = output_dir
        

class BugFixManagerTestResult(WorkerTestResult):

    config_file = Config.get_default_config_file(__file__)

    def execute_auto_task_one_input(self):
        # STEP1 Configure the input output
        self.config.RETURN is None
        input_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_input"], self.config.INIT)
        output_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_output"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        actual_dir = os.path.join(self.config.SRV_BASE_DIR + self.config.the_root_dict["test_data_actual"],
                                  self.config.METHOD,
                                  self.config.OUTPUT)
        config = Config()
        tc_output_dir = os.path.join(self.config.SRV_BASE_DIR, "test", "Output", self.config.ClassName, self.config.TestID)
        config.update_config(InputDir=input_dir)
        config.update_config(OutputDir=tc_output_dir)
        config.update_config(input_dir=input_dir)
        config.refresh_config()
        Folder.create_directory(actual_dir, exist_del=True)
        # STEP2 run the method
        self.capture.start()
        self.config.TestClassObject = self.config.MyClass(config)
        copied_config = getattr(self.config.TestClassObject, "get_config")()
        self.config.RETURN = None
        if isinstance(self.config.INPUT, list) and self.config.INPUT == []:
            getattr(self.config.TestClassObject, self.config.METHOD)()
        elif isinstance(self.config.INPUT, str):
            getattr(self.config.TestClassObject, self.config.METHOD)(self.config.INPUT)
        # copy data from session to actual result
        file_lists = Folder.scan_all_files(copied_config.StepOutPutDir[copied_config.StepList[-1]])
        for file in file_lists:
            src = file
            dst = os.path.join(actual_dir, os.path.basename(file))
            copyfile(src, dst)
        folder_lists = Folder.scan_all_folder(copied_config.StepOutPutDir[copied_config.StepList[-1]])
        for folder in folder_lists:
            src = folder
            dst = os.path.join(actual_dir, os.path.basename(folder))
            copytree(src, dst)
        self.capture.stop()
        self.update_error_warn_info()
        self.config.ACTUAL = actual_dir
        self.config.OUTPUT = output_dir

