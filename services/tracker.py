# encoding: utf-8
import os
import sys
import copy
import docx.package
import docx.parts.document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from settings.config import Config
from services.wordML import WordML


class Tracker:
    """Tracker"""
    __count = 0
    __default_template_docx = os.path.join(Config.BASE_DIR, "services", "Templates","Template_TrackerReport.docx")

    def __init__(self):
        package = docx.package.Package.open(self.__default_template_docx)
        main_document_part = package.main_document_part
        assert isinstance(main_document_part, docx.parts.document.DocumentPart)
        self.document = main_document_part.document
        self.document.add_heading('Auto Format Report', 1)
        para0 = self.document.paragraphs[0]
        self.document.element.body.remove(para0._p)
        self.__content_modify_dict = {}
        self.__count_modify_dict = {}
        self.__total = 0

    def add_modify_items(self, auto_task_name, before, after):
        self.__total += 1
        if auto_task_name not in self.__count_modify_dict:
            self.__count_modify_dict[auto_task_name] = 1
            self.__content_modify_dict[auto_task_name] = [[before, after]]
        else:
            self.__count_modify_dict[auto_task_name] += 1
            self.__content_modify_dict[auto_task_name].append([before, after])

    def create_summary_report(self):
        print("      creating summary report...")
        self.document.add_heading('Number of Changes', 2)
        self.document.add_paragraph("Total of changes: %s" % self.__total)
        self.document.add_paragraph("Table 1: Number of Changes by Auto Task Name")
        table = self.document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id'
        hdr_cells[1].text = 'Auto Task Name'
        hdr_cells[2].text = 'Number of Changes'
        for i, auto_task_name in enumerate(self.__count_modify_dict):
            print("      - %s : %s" % (auto_task_name,  self.__count_modify_dict[auto_task_name]))
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = auto_task_name
            row_cells[2].text = str(self.__count_modify_dict[auto_task_name])

    def create_detail_report(self):
        print("      creating detail report...")
        self.document.add_heading('Track Change Description', 2)
        idx = 1
        for auto_task_name in self.__content_modify_dict:
            self.document.add_heading(auto_task_name, 3)
            for [before, after] in self.__content_modify_dict[auto_task_name]:
                change_key = auto_task_name + " No." + "0" * (3 - len("%s" % idx)) + "%d" % idx
                print("      - %s: %s / %s" % (auto_task_name, idx, self.__total))
                self.document.add_heading(change_key, 4)
                added_before = self.document.add_paragraph("Before : ")
                added_after = self.document.add_paragraph("After  : ")
                if before is not None and after is not None:
                    if isinstance(before, str):
                        added_before.text = "Before : " + before
                        added_after.text = "After : " + after
                    elif isinstance(before, Paragraph):
                        if WordML.is_heading(before):
                            before = WordML.remove_style(before)
                        pi = WordML.copy_paragraph_after(before, added_before)
                        if WordML.is_heading(after):
                            after = WordML.remove_style(after)
                        pi = WordML.copy_paragraph_after(after, added_after)

                elif before is not None and after is None:
                    if isinstance(before, str):
                        added_before.text = "Before : " + before
                        added_after.text = "After : " + "<del>"
                    elif isinstance(before, Paragraph):
                        if WordML.is_heading(before):
                            before = WordML.remove_style(before)
                        pi = WordML.copy_paragraph_after(before, added_before)
                        after = WordML.highlight_deleted_paragraph(before)
                        pi = WordML.copy_paragraph_after(after, added_after)
                    elif isinstance(before, Table):
                        before = WordML.highlight_delete_table(before)
                        # table = self.document.add_table(rows=1, cols=17)
                        WordML.copy_table_after(before, added_after)
                        # self.document.element.body.replace(table._tbl, copy.deepcopy(before._tbl))
                idx += 1

    def save(self, filename):
        print("      save to %s..." % filename)
        self.document.save(filename)

    def add_summary_modify_items(self, auto_task_name, count):
        self.__total += count
        if auto_task_name not in self.__count_modify_dict:
            self.__count_modify_dict[auto_task_name] = count
            # self.__content_modify_dict[auto_task_name] = [[before, after]]
        else:
            self.__count_modify_dict[auto_task_name] += count
            # self.__content_modify_dict[auto_task_name].append([before, after])