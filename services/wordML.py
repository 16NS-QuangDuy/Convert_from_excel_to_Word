# encoding: utf-8
import re
import os
import copy
from copy import deepcopy
import yaml
import docx.package
import docx.parts.document
import docx.parts.numbering
import docx.parts.styles
from docx import Document
from docx.table import _Cell, Table
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from settings.config import Config
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'
BODY = WORD_NAMESPACE + 'body'
TBL = WORD_NAMESPACE + 'tbl'
import win32com.client
import time
from services.elmHelper import ElmHelper
import os


class WordML:
    """WordML"""
    auto_worker_name = "WordML"
    config_file = Config.get_default_config_file(the_script_file=__file__,the_config_name="wordML_config.yaml")

    def __init__(self, filename=None, package=None):
        self.auto_worker_name = self.__class__.__name__
        self.filename = filename
        self.config = Config()
        self.package = package
        if filename is not None:
            self.package = docx.package.Package.open(filename)
        if self.package is not None:
            self.main_document_part = self.package.main_document_part
            assert isinstance(self.main_document_part, docx.parts.document.DocumentPart)
            self.document = self.main_document_part.document
            self.related_parts = self.document.part.related_parts
            try:
                self.numbering_part = self.main_document_part.numbering_part
                assert isinstance(self.numbering_part, docx.parts.numbering.NumberingPart)
            except:
                self.numbering_part = None
                pass            
        self.blocks = []
        self.layout = []
        self.layout_all_contents = []
        self.dict_numstr_headings = {}
        self.heading_numbering_dict = {}
        self.heading_list = []
        self.styles = []
        self.number_heading_dict = dict()
        self.number_heading_list = []
        self.number_list_paragraph_dict = dict()
        self.layout_images = []
        self.layout_list_paragraph_numbering = dict()
        self.list_paragraph_numbering_dict = dict()
        self.level1_style = "Heading1"
        self.StyleNumIdDict = self.create_styleid_numid_dict()
        Config.set_attr_from_yaml(self.__dict__["config"], self.config_file, self.auto_worker_name)
        self.BLOCKS = self.blocks

    def iter_block_items(self, parent=None):
        if parent == None:
            parent = self.document
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def iter_block_items_external(parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)


    @staticmethod
    def iter_paragraphs(parent, recursive=True):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
        elif isinstance(parent, docx.table._Cell):
            parent_elm = parent._tc
        else:
            raise TypeError(repr(type(parent)))

        for child in parent_elm.iterchildren():
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                yield docx.text.paragraph.Paragraph(child, parent)
            elif isinstance(child, docx.oxml.table.CT_Tbl):
                if recursive:
                    table = docx.table.Table(child, parent)
                    for row in table.rows:
                        try:
                            for cell in row.cells:
                                for child_paragraph in WordML.iter_paragraphs(cell):
                                    yield child_paragraph
                        except:
                            print("        Error Table")
                            pass

    def get_all_blocks(self):
        if len(self.blocks) <= 0:
            for i, block in enumerate(self.iter_block_items()):
                self.blocks.append(block)
        return self.blocks

    @staticmethod
    def get_full_latest_run_paragraph(p):
        r_lst = []
        temp_r_list = p._p.r_lst
        for child in p._p.getchildren():
            if child.tag == (WORD_NAMESPACE + "ins"):
                ins_runs = child.findall(".//w:r", namespaces=child.nsmap)
                r_lst.extend(ins_runs)
            if child.tag == (WORD_NAMESPACE + "r"):
                if temp_r_list.count(child) >= 1:
                    r_lst.append(child)
        return r_lst

    @staticmethod
    def get_full_latest_paragraph_text(p):
        paragraph_runs = WordML.get_full_latest_run_paragraph(p)
        text = ""
        for run in paragraph_runs:
            temp_t_list = run.t_lst
            t_lst = run.findall('.//w:t', namespaces=run.nsmap)  ##ONLY DIRECT CHILD
            for t in t_lst:
                if temp_t_list.count(t) >= 1:
                    text = text + t.text
        return text

    @staticmethod
    def get_full_old_run_paragraph(p):
        r_lst = []
        temp_r_list = p._p.r_lst
        for child in p._p.getchildren():
            if child.tag == (WORD_NAMESPACE + "del"):
                del_runs = child.findall(".//w:r", namespaces=child.nsmap)
                r_lst.extend(del_runs)
            if child.tag == (WORD_NAMESPACE + "r"):
                if temp_r_list.count(child) >= 1:
                    r_lst.append(child)
        return r_lst

    @staticmethod
    def get_full_old_paragraph_text(p):
        paragraph_runs = WordML.get_full_old_run_paragraph(p)
        text = ""
        for run in paragraph_runs:
            t_lst = run.findall('.//w:t', namespaces=run.nsmap)
            t_lst_del = run.findall('.//w:delText', namespaces=run.nsmap)
            if len(t_lst) > 0:
                temp_t_list = run.t_lst
                for t in t_lst:
                    if temp_t_list.count(t) >= 1:
                        text = text + t.text
            if len(t_lst_del) > 0:
                for t in t_lst_del:
                    text = text + t.text
        return text

    @staticmethod
    def delete_paragraph(paragraph):
        try:
            parent = paragraph._element.getparent()
            # cell must have at least 1 paragraph
            if parent.tag == WORD_NAMESPACE + 'tc':
                p_list = [child for child in parent.getchildren() if child.tag == WORD_NAMESPACE + 'p']
                if len(p_list) > 1:
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None
                else:
                    paragraph.text = ""
            else:
                if ElmHelper.finditer3(paragraph._element, tag_list=["headerReference", "footerReference"], ns="w"):
                    pass
                else:
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._element = None
        except:
            print("ERROR: delete_paragraph...")
            pass

    @staticmethod
    def delete_block(block):
        try:
            WordML.delete_paragraph(block)
        except:
            print ("ERROR: delete_block...")
            pass

    @staticmethod
    def remove_prefix_numbering(text):
        text = text.strip()
        match = re.match(r"^([\d+\.]+\d+)(.*)", text)
        if match:
            text = match.group(2)
            text = text.strip()
        return text

    @staticmethod
    def get_heading(block):
        pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
        if pStyle is not None:
            heading = block.style.name
            if heading.startswith("Heading"):
                return heading.replace(" ", "")
        return None

    @staticmethod
    def get_style_name(block):
        pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
        if pStyle is not None:
            return block.style.name
        return None

    @staticmethod
    def is_heading(block):
        pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
        if pStyle is not None:
            heading = block.style.name
            if heading.startswith("Heading"):
                return True
        return False

    @staticmethod
    def remove_style(block):
        pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
        if pStyle is not None:
            if (pStyle.val.startswith("Heading") or pStyle.val in ["1", "2", "3", "4", "5", "6", "7", "8",
                                                                   "9"]):
                p = pStyle.getparent()
                p.remove(pStyle)
        return block

    @staticmethod
    def is_hidden_paragraph(block):
        paragraph = block
        for run in paragraph.runs:
            run = run._r
            rPr = run.find('.//w:rPr', namespaces=run.nsmap)
            if rPr is None:
                continue
            vanish = rPr.find('.//w:vanish', namespaces=rPr.nsmap)
            if vanish is not None:
                return True
        return False

    @staticmethod
    def make_run_hidden_and_shadow(run, ndisclosed_dict=None):
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        if ndisclosed_dict is not None:
            rgb = ndisclosed_dict["TextColor"]
            WordML.color_run(run, RGBColor(rgb[0], rgb[1], rgb[2]))
            shading_elm_1 = parse_xml(r'<w:shd {} w:val="pct15" w:color="auto" w:fill="FFFFFF"/>'.format(nsdecls('w')))
            shading_elm_2 = parse_xml(r'<w:vanish {}/>'.format(nsdecls('w')))
            run = run._r
            rPr = run.find('.//w:rPr', namespaces=run.nsmap)
            if rPr is not None:
                if "ShadowColor" in ndisclosed_dict:
                    if ndisclosed_dict["ShadowColor"] != "auto":
                        shading_elm_1 = parse_xml((
                            r'<w:shd {} w:val="pct15" w:color="%s" w:fill="FFFFFF"/>' % (ndisclosed_dict["ShadowColor"])).format(nsdecls('w')))
                if ndisclosed_dict["MadeShadow"]:
                    rPr.append(shading_elm_1)
                if ndisclosed_dict["MadeHidden"]:
                    rPr.append(shading_elm_2)

    @staticmethod
    def make_paragraph_hidden_and_shadow(paragraph, ndisclosed_dict=None):
        for run in paragraph.runs:
            WordML.make_run_hidden_and_shadow(run, ndisclosed_dict)
        try:
            if ndisclosed_dict["TextColor"]:
                rgb = ndisclosed_dict["TextColor"]
                WordML.color_paragraph(paragraph, RGBColor(rgb[0], rgb[1], rgb[2]))
        except:
            pass

    @staticmethod
    def make_cells_hidden_and_shadow(cells, ndisclosed_dict=None):
        for cell in cells:
            for paragraph in cell.paragraphs:
                WordML.make_paragraph_hidden_and_shadow(paragraph, ndisclosed_dict)

    @staticmethod
    def make_block_hidden_and_shadow(block, ndisclosed_dict=None):
        if isinstance(block, Table):
            for row in block.rows:
                WordML.make_cells_hidden_and_shadow(row.cells, ndisclosed_dict)
        elif isinstance(block, Paragraph):
            WordML.make_paragraph_hidden_and_shadow(block, ndisclosed_dict)

    @staticmethod
    def make_sub_paragraph_hidden_and_shadow(paragraph, subparagraph):
        run1 = paragraph.add_run()
        run2 = paragraph.add_run()
        run3 = paragraph.add_run()
        text = paragraph.text
        result = text.find(subparagraph)
        paragraph.text = ""
        if result != -1:
            run1.text = text[:result]
            run2.text = subparagraph
            run3.text = text[result+ len(subparagraph):]
        WordML.make_run_hidden_and_shadow(run2._r)

    @staticmethod
    def color_paragraph(paragraph, color):
        color_elm = OxmlElement("w:color")
        color_elm.val = color
        p = paragraph._p
        rPr = p.find('.//w:pPr/w:rPr', namespaces=p.nsmap)
        while True:
            color1 = p.find('.//w:pPr/w:rPr/w:color', namespaces=p.nsmap)
            if color1 is not None:
                p = color1.getparent()
                p.remove(color1)
            else:
                break
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            pPr = p.find('.//w:pPr', namespaces=p.nsmap)
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.append(pPr)
                pPr = p.find('.//w:pPr', namespaces=p.nsmap)
            pPr.append(rPr)
            rPr = p.find('.//w:pPr/w:rPr', namespaces=p.nsmap)
        rPr.append(color_elm)

    @staticmethod
    def color_run(run, color):
        font = run.font
        font.color.rgb = color
        # rPr = run.find('.//w:rPr', namespaces=run.nsmap)
        # color1 = run.find('.//w:rPr/w:color', namespaces=run.nsmap)
        # color_elm = OxmlElement("w:color")
        # color_elm.val = color
        # if color1 is not None:
        #     p = color1.getparent()
        #     p.remove(color1)
        # if rPr is None:
        #     rPr = OxmlElement('w:rPr')
        #     run.append(rPr)
        #     rPr = run.find('.//w:rPr', namespaces=run.nsmap)
        # rPr.append(color_elm)

    @staticmethod
    def is_cjk(character):
        """"
        Checks whether character is CJK.
            >>> is_cjk(u'\u33fe')
            True
            >>> is_cjk(u'\uFE5F')
            False
        :param character: The character that needs to be checked.
        :type character: char
        :return: bool
        """
        return any([start <= ord(character) <= end for start, end in
                    [(4352, 4607), (11904, 42191), (43072, 43135), (44032, 55215),
                     (63744, 64255), (65072, 65103), (65381, 65500),
                     (131072, 196607)]
                    ])

    @staticmethod
    def is_cjk_strings(string):
        i = 0
        count = 0
        while i<len(string):
            if WordML.is_cjk(string[i]):
                count +=1
            i += 1
        if count >=1:
            return True
        else:
            return False

    @staticmethod
    def highlight_cell_keep(cell):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="7cfc00"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def highlight_cell_delete(self, cell):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="dcdcdc"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)
    
    def highlight_cell_error(self, cell):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="ff1493"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)


    @staticmethod
    def highlight_deleted_paragraph(paragraph):
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
        return paragraph

    @staticmethod
    def highlight_deleted_runs(runs):
        shading_elm_1 = parse_xml(r'<w:shd {} w:val="pct15" w:color="auto" w:fill="cccccc"/>'.format(nsdecls('w')))
        for run in runs:
            rPr = run.find('.//w:rPr', namespaces=run.nsmap)
            if rPr is not None:
                rPr.append(shading_elm_1)

    @staticmethod
    def highlight_run_delete(self, run):
        shading_elm_1 = parse_xml(r'<w:shd {} w:val="pct15" w:color="auto" w:fill="cccccc"/>'.format(nsdecls('w')))
        rPr = run.find('.//w:rPr', namespaces=run.nsmap)
        if rPr is not None:
            rPr.append(shading_elm_1)

    @staticmethod
    def highlight_delete_table(block):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="dcdcdc"/>'.format(nsdecls('w')))
        for column_index in range(0, len(block.columns)):
            for row_index in range(0, len(block.rows)):
                cell = block.cell(row_index, column_index)
                cell._tc.get_or_add_tcPr().append(shading_elm_1)
        return block

    def get_layout_of_headings(self):
        """ Headings and its information into dict
            Iter throught block (self.iter_block_items) and find pStyle element
        Params:: 
            self.config_file
        Returns::
            self.layout = {
                "block_id": ,
                "text": ,
                "level": ,
                "style: ,
                "heading": ,
            }
        """
        auto_task_name = WordML.auto_worker_name
        __target_spec_config_dict = {}
        level_dict = {}
        with open(self.config_file, 'r', encoding="utf-8") as stream:
            __target_spec_config_dict = yaml.safe_load(stream)
        if auto_task_name in __target_spec_config_dict:
            level_dict = __target_spec_config_dict[auto_task_name]["LevelDict"]

        for j, block in enumerate(self.iter_block_items()):
            if isinstance(block, Paragraph):
                if block.text.strip() == "":
                    continue
                pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
                if pStyle is not None:
                    style = pStyle.val
                    heading = block.style.name
                    heading = heading.replace(" ", "")
                    if heading.startswith("Heading"):
                        if level_dict is not None:
                            for level in level_dict:
                                if heading in level_dict[level]:
                                    self.layout.append({"block_id": j, "text": block.text.strip(), "level": level, "style": style, "heading": heading})
                                    found = True
                                    break
        return self.layout

    def generate_all_heading_numbering(self):
        """ self.layout added numbering
        """
        heading_list = []
        self.level1_style = "Heading1"
        for i,  layout_i in enumerate(self.layout):
            style = layout_i["style"]
            level = layout_i["level"]
            if level =="level1" and style != self.level1_style:
                self.level1_style = style
                break
        for i,  layout_i in enumerate(self.layout):
            num = self.generate_heading_numbering(i, self.level1_style)
            text = layout_i["text"]
            style = layout_i["style"]
            level = layout_i["level"]
            block_id = layout_i["block_id"]
            heading = layout_i["heading"]
            if self.is_prefix_numbering_heading(text):
                text = text.replace(num, "")
                text = text.strip()
            heading_list.append({"num": num, "text": text, "level": level, "style": style, "heading": heading, "block_id": block_id})
        self.number_heading_list = heading_list
        return heading_list

    def generate_all_number_heading_dict(self, key="text"):
        """ return a dict of numbering of heading with key is block_id
        """
        if self.numbering_part is None:
            return self.number_heading_dict
        auto_task_name = WordML.auto_worker_name
        __target_spec_config_dict = {}
        with open(self.config_file, 'r', encoding="utf-8") as stream:
            __target_spec_config_dict = yaml.safe_load(stream)
        if auto_task_name in __target_spec_config_dict:
            level_dict = __target_spec_config_dict[auto_task_name]["LevelDict"]
            self.numbering = __target_spec_config_dict[auto_task_name]
            self.indent = self.numbering["IndentPrefix"]
            self.level_list =self.numbering["LevelList"]
            self.preserve = self.numbering["PreservedText"]
            self.level_indent_dict = self.numbering["LevelIndentDict"]
            self.numbering_level_list = self.numbering["NumberingLevelList"]
        number_heading_list = {}
        self.level1_style = "Heading1"
        for i,  layout_i in enumerate(self.layout):
            style = layout_i["style"]
            level = layout_i["level"]
            if level =="level1" and style != self.level1_style:
                self.level1_style = style
                break
        for i,  layout_i in enumerate(self.layout):
            num = self.generate_heading_numbering(i, self.level1_style)
            text = layout_i["text"]
            style = layout_i["style"]
            level = layout_i["level"]
            block_id = layout_i["block_id"]
            heading = layout_i["heading"]
            if self.is_prefix_numbering_heading(text):
                text = text.replace(num, "")
                text = text.strip()
            text = "%s%s%s" % (num, self.preserve, text)
            if key == "text":
                number_heading_list[text] = {"num": num, "text": text, "level": level, "style": style, "heading": heading, "block_id": block_id}
            if key == "block_id":
                number_heading_list["%s" % block_id] = {"num": num, "text": text, "level": level, "style": style,
                                                        "heading": heading, "block_id": block_id}
        self.number_heading_dict = number_heading_list
        return number_heading_list

    def generate_heading_numbering(self, i, level1_style="Heading1"):
        """ from level, start and lvltext => calculate actual number of numbering
        """
        level_dict = {
            "level1": "1",
            "level2": "2",
            "level3": "3",
            "level4": "4",
            "level5": "5",
            "level6": "6",
            "level7": "7",
            "level8": "8",
            "level9": "9",
            "levelA": "A",
        }
        key_text = ""
        level_list = []
        count_dict = {}
        enable_count_dict = {}
        for level in level_dict:
            count_dict[level] = 0
            enable_count_dict[level] = True
        cur_level_dict = self.layout[i]
        if self.is_prefix_numbering_heading(cur_level_dict["text"]):
            realnum =self.get_prefix_numbering(cur_level_dict["text"])
            return realnum
        level_list.append(cur_level_dict["level"])
        for j in range(i, -1, -1):
            j_cur_level_dict = self.layout[j]
            if j_cur_level_dict["level"] == cur_level_dict["level"]:
                if enable_count_dict[j_cur_level_dict["level"]]:
                    if not self.is_prefix_numbering_heading(j_cur_level_dict["text"]):
                        count_dict[j_cur_level_dict["level"]] += 1
                    continue
            if j_cur_level_dict["level"] < cur_level_dict["level"]:
                level_list.insert(0, j_cur_level_dict["level"])
                if not self.is_prefix_numbering_heading(j_cur_level_dict["text"]):
                    count_dict[j_cur_level_dict["level"]] += 1
                enable_count_dict[cur_level_dict["level"]] = False
                cur_level_dict = j_cur_level_dict
                continue
            if j_cur_level_dict["level"] == "level1":
                break
        cur_level_dict = self.layout[i]
        heading = cur_level_dict["style"]
        if heading not in self.heading_numbering_dict:
            realnum = ""
            return realnum
        abstractnum = self.heading_numbering_dict[heading]
        lvlText = abstractnum["lvlText"]
        realnum = lvlText
        if "(%5)" in lvlText:
            count = count_dict["level5"]
            level_int =5
            abstracttext = "%" + "%d" % level_int
            realnum = realnum.replace(abstracttext, "%s" % count)
            return realnum
        else:
            for i, level in enumerate(level_list):
                if level == "level1":
                    abstractnum = self.heading_numbering_dict[level1_style]
                    start = abstractnum["Start"]
                    level_int = i + 1
                    abstracttext = "%"+"%d" % level_int
                    realnum = realnum.replace(abstracttext, "%s" % start)
                else:
                    count = count_dict[level]
                    level_int = i + 1
                    abstracttext = "%"+"%d" % level_int
                    realnum = realnum.replace(abstracttext, "%s" % count)
            return realnum

    @staticmethod
    def is_prefix_numbering_heading(text):
        text = text.strip()
        match01 = re.match(r"^(\d+\.)(.*)", text)
        match02 = re.match(r"^(\d+[A-Z]+\.)(.*)", text)
        if match01 or match02:
            return True
        else:
            return False

    @staticmethod
    def get_prefix_numbering(text):
        text = text.strip()
        heading_num = ""
        arr = text.split(".")
        for num in arr[:len(arr)-1]:
            heading_num += "%s." % num
        last_num = arr[len(arr)-1]
        match01 = re.match(r"^(\d+)(.*)", last_num) # case 87.5A.1
        match02 = re.match(r"^(\d+[A-Z]*)\s+(.*)", last_num) # case 87.5A
        if match02:
            last_num = "%s" % match02.group(1)
            heading_num += last_num
            return heading_num
        if match01:
            last_num = "%s" % match01.group(1)
            heading_num += last_num
            return heading_num
        return heading_num

    def get_all_heading_numbering_dict(self):
        if self.numbering_part is None: return
        abstractNumlist = self.numbering_part._element.findall(".//w:abstractNum",
                                                               namespaces=self.numbering_part._element.nsmap)
        for abstractNum in abstractNumlist:
            list_lvl = abstractNum.findall(".//w:lvl", namespaces=abstractNum.nsmap)
            for lvl in list_lvl:
                pStyle = lvl.find(WORD_NAMESPACE + 'pStyle')
                if pStyle is not None:
                    style = pStyle.val
                    ilvl = lvl.attrib[WORD_NAMESPACE + 'ilvl']
                    start = lvl.find(WORD_NAMESPACE + 'start')
                    if start is not None:
                        start = start.attrib[WORD_NAMESPACE + 'val']
                    numFmt = lvl.find(WORD_NAMESPACE + 'numFmt')
                    if numFmt is not None:
                        numFmt = numFmt.attrib[WORD_NAMESPACE + 'val']
                    lvlText = lvl.find(WORD_NAMESPACE + 'lvlText')
                    if lvlText is not None:
                        lvlText = lvlText.attrib[WORD_NAMESPACE + 'val']
                    self.heading_numbering_dict[style] = {"Style": style, "ilvl": ilvl, "numFmt": numFmt, "lvlText": lvlText, "Start": start}

        #self.layout.append({"block_id": j, "text": block.text.strip(), "level": level, "style": style, "heading": heading})
        if len(self.blocks)<= 0:
            self.get_all_blocks()
        for li in self.layout:
            try:
                style = li["style"]
                block_id =li["block_id"]
                block = self.blocks[block_id]
                lvl = self.get_paragraph_lvl(block)
                if lvl is not None:
                    style = style
                    ilvl = lvl.attrib[WORD_NAMESPACE + 'ilvl']
                    start = lvl.find(WORD_NAMESPACE + 'start')
                    if start is not None:
                        start = start.attrib[WORD_NAMESPACE + 'val']
                    numFmt = lvl.find(WORD_NAMESPACE + 'numFmt')
                    if numFmt is not None:
                        numFmt = numFmt.attrib[WORD_NAMESPACE + 'val']
                    lvlText = lvl.find(WORD_NAMESPACE + 'lvlText')
                    if lvlText is not None:
                        lvlText = lvlText.attrib[WORD_NAMESPACE + 'val']
                    self.heading_numbering_dict[style] = {"Style": style, "ilvl": ilvl, "numFmt": numFmt, "lvlText": lvlText,
                                                          "Start": start}
            except Exception as e:
                print ( "Error: get_all_heading_numbering_dict" + str(e))
                pass
        if self.heading_numbering_dict.get('Heading1', None) is None and self.heading_numbering_dict.get('1', None) is None :
            for li in self.layout:
                style = li["style"]
                heading = li["heading"]
                block_id = li["block_id"]
                text = li["text"]
                block = self.blocks[block_id]
                ilvl = 0
                if heading == "Heading1":
                    if text.startswith("Section"):
                        numFmt = "decimal"
                        lvlText = "Section %1"
                        text = text.replace("Section", "").strip()
                        start = text.split(" ")[0]
                    else:
                        numFmt = "decimal"
                        text = text.strip()
                        lvlText = text.split(" ")[0]
                        match = re.search(r"[U]*[0-9]+", text)
                        if match:
                            start = match.group(0)
                    # print("        This chapter has no Heading 1 numbering")
                    self.heading_numbering_dict[style] = {"Style": style, "ilvl": ilvl, "numFmt": numFmt, "lvlText": lvlText,
                                                          "Start": start}
        return self.heading_numbering_dict

    def change_color_all_abstract_num(self, rgblist):
        abstractNumlist = self.numbering_part._element.findall(".//w:abstractNum",
                                                               namespaces=self.numbering_part._element.nsmap)
        shading_elm = parse_xml(
            r'<w:color {} w:val="auto"/>'.format(nsdecls('w')))
        for abstractNum in abstractNumlist:
            list_lvl = abstractNum.findall(".//w:lvl", namespaces=abstractNum.nsmap)
            for lvl in list_lvl:
                is_removed_already = False
                pStyle = lvl.find(WORD_NAMESPACE + 'pStyle')
                if pStyle is not None and pStyle.val.startswith("Heading"):
                    rPr = lvl.find(WORD_NAMESPACE + 'rPr')
                    if rPr is not None:
                        _color = rPr.find(WORD_NAMESPACE + 'color')
                        if _color is not None:
                            color = _color.attrib[WORD_NAMESPACE + 'val']
                            if isinstance(color, str) and color != 'auto':
                                i = int(color, 16)
                                f = format(i, "06x")
                                (R, G, B) = (f[:2], f[2:4], f[4:6])
                                color = RGBColor(int(R, 16), int(G, 16), int(B, 16))
                            if isinstance(color, RGBColor):
                                if color in rgblist:
                                    p = _color.getparent()
                                    p.remove(_color)
                                    is_removed_already = True
                    if is_removed_already:
                        lvl.append(shading_elm)

    @staticmethod
    def is_table_block(block):
        return isinstance(block, Table)

    @staticmethod
    def is_paragraph_block(block):
        return isinstance(block, Paragraph)

    @staticmethod
    def is_empty_line(row):
        count = 0
        cells_text = ElmHelper.get_tr_text(row._tr)
        is_empty_line = cells_text.count("") == len(cells_text)
        if len(ElmHelper.finditer(row._tr, "drawing")) > 0 or \
            len(ElmHelper.finditer(row._tr, "pic", "pic")) > 0:
            is_empty_line = False
        return is_empty_line

    def get_layout_of_target_spec(self):
        auto_task_name = "WordML"
        __target_spec_config_dict = {}
        level_dict = {}
        with open(self.config_file, 'r', encoding="utf-8") as stream:
            __target_spec_config_dict = yaml.safe_load(stream)
        if auto_task_name in __target_spec_config_dict:
            level_dict = __target_spec_config_dict[auto_task_name]["LevelDict"]
        self.heading_block_idx_list = []
        for j, block in enumerate(self.iter_block_items()):
            if isinstance(block, Paragraph):
                pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
                if pStyle is not None:
                    style = pStyle.val
                    heading = block.style.name
                    found = False
                    if heading.startswith("Heading"):
                        if level_dict is not None:
                            for level in level_dict:
                                if heading in level_dict[level]:
                                    self.layout_all_contents.append({"block_id": j, "text": block.text.strip(), "level": level, "style": style, "heading": heading})
                                    self.layout.append({"block_id": j, "text": block.text.strip(), "level": level, "style": style, "heading": heading})
                                    self.heading_block_idx_list.append(j)
                                    found = True
                                    break
                    if not found:
                        self.layout_all_contents.append(
                            {"block_id": j, "text": block.text.strip(), "level": "levelA", "style": style, "heading": heading})
                else:
                    self.layout_all_contents.append(
                        {"block_id": j, "text": block.text.strip(), "level": "levelA", "style": None, "heading": None})
            else:
                self.layout_all_contents.append(
                {"block_id": j, "text": "<<table>>", "level": "levelA", "style": None, "heading": None})
        return self.layout_all_contents

    def get_all_tables_with_title_and_subsection(self):
        self.get_all_heading_numbering_dict()
        self.heading_list = self.generate_all_heading_numbering()
        self.table_list = []
        record = []
        for (i, block) in enumerate(self.iter_block_items()):
            if isinstance(block, Table):
                record = self.parse_table_into_records(block)
                table_title = self.get_table_title(i)
                subsection = self.get_heading_of_block(i)
                self.table_list.append({"record": record, "table_title": table_title, "subsection": subsection, "block_id": i})
        return self.table_list

    def generate_table_title_dict(self):
        result_dict = dict()
        for (i, block) in enumerate(self.iter_block_items()):
            if isinstance(block, Table):
                table_title = self.get_table_title(i)
                result_dict["%s" % i] = {"table_title": table_title}
        return result_dict

    def get_table_title(self, j):
        is_title_found = False
        bound = 5
        i = j
        while not is_title_found:
            i = i - 1
            bound = bound - 1
            if i < 0 or bound < 0:
                break
            try:
                layout_i = self.layout_all_contents[i]
            except IndexError:
                print ("Error: ", i )
                return ""
            text = layout_i["text"]
            if text.startswith("Table"):
                is_title_found = True
                return text
            if layout_i["heading"] == "heading":
                is_title_found = False
                return ""

    def get_heading_of_block(self, j):
        is_heading_found = False
        i = len(self.heading_list)
        while not is_heading_found:
            i = i - 1
            if i < 0:
                break
            layout_i = self.heading_list[i]
            block_id = layout_i["block_id"]
            if block_id < j:
                is_heading_found = True
                return layout_i

    def parse_table_into_records(self, block):
        records = []
        for int_idx, row in enumerate(block.rows):
            rc_dict = []
            if self.is_empty_line(row):
                continue
            for fld_idx in range(len(row.cells)):
                cell = row.cells[fld_idx]
                field = cell.text.strip().replace("\n", "")
                rc_dict.append(field)
            records.append(copy.deepcopy(rc_dict))
        return records

    @staticmethod
    def copy_first_table_from_file(docx_file):
        package = docx.package.Package.open(docx_file)
        main_document_part = package.main_document_part
        assert isinstance(main_document_part, docx.parts.document.DocumentPart)
        document = main_document_part.document
        tbl_bit_description = copy.deepcopy(document.tables[0])
        copied_tbl = copy.deepcopy(tbl_bit_description._tbl)
        return copied_tbl

    @staticmethod
    def copy_table_from_file(docx_file, idx=0):
        package = docx.package.Package.open(docx_file)
        main_document_part = package.main_document_part
        assert isinstance(main_document_part, docx.parts.document.DocumentPart)
        document = main_document_part.document
        tbl_bit_description = copy.deepcopy(document.tables[idx])
        # copied_tbl = copy.deepcopy(tbl_bit_description._tbl)
        return tbl_bit_description

    @staticmethod
    def copy_textbox_pict_from_file(docx_file,idx=0):
        package = docx.package.Package.open(docx_file)
        main_document_part = package.main_document_part
        assert isinstance(main_document_part, docx.parts.document.DocumentPart)
        document = main_document_part.document
        textbox_pict = copy.deepcopy(document.paragraphs[idx])
        return textbox_pict

    @staticmethod
    def modify_special_text(run):
        """
        Modify run style
        """
        rFonts_elm = parse_xml(r'<w:rFonts {} w:ascii="Arial" w:eastAsia="Arial" w:hAnsi="Arial"/>'.format(nsdecls('w')))
        if "¾" in run.text:
            rFonts_elm = parse_xml(
                r'<w:rFonts {} w:ascii="Symbol" w:eastAsia="Symbol" w:hAnsi="Symbol"/>'.format(nsdecls('w')))
        sz_elm = parse_xml(r'<w:sz {} w:val="14"/>'.format(nsdecls('w')))
        rPr_elm = parse_xml(r'<w:rPr {}/>'.format(nsdecls('w')))
        rPr = run.find('.//w:rPr', namespaces=run.nsmap)
        if rPr is None:
            rPr = run.insert(0, rPr_elm)
        rPr = run.find('.//w:rPr', namespaces=run.nsmap)
        rFonts = run.find('.//w:rPr/w:rFonts', namespaces=run.nsmap)
        if rFonts is not None:
            rPr.remove(rFonts)
        rPr.append(rFonts_elm)
        sz = run.find('.//w:rPr/w:sz', namespaces=run.nsmap)
        if sz is not None:
            rPr.remove(sz)
        rPr.append(sz_elm)

    @staticmethod
    def set_cell_text_for_bit_name(cell, newtext, alignment_val=None):
        cell.text = newtext
        if alignment_val is not None:
            for i, para in enumerate(cell.paragraphs):
                para.alignment = alignment_val

    @staticmethod
    def set_cell_text(cell, newtext, alignment_val=None):
        for i, para in enumerate(cell.paragraphs):
            if len(para.runs) == 0:
                para.add_run(newtext)
            for k, run in enumerate(para.runs):
                if k == 0:
                    run.text = newtext
                else:
                    run.text = ""
        if alignment_val is not None:
            for i, para in enumerate(cell.paragraphs):
                para.alignment = alignment_val

    @staticmethod
    def set_cell_font_size(cell, name="Arial", size=9):
        for i, para in enumerate(cell.paragraphs):
            for k, run in enumerate(para.runs):
                font = run.font
                font.size = Pt(size)
                font.name = name

    @staticmethod
    def get_merge_val(row, column_index):
        tr = row._tr
        alltc = tr.findall(".//w:tc", namespaces = tr.nsmap)
        if alltc is not None and len(alltc) >0 and column_index <len(alltc) :
            tc = alltc[column_index]
            tcPr = tc.find(".//w:tcPr", namespaces = tc.nsmap)
            if tcPr is not None:
                vMerge = tcPr.find(".//w:vMerge", namespaces = tcPr.nsmap)
                if vMerge is not None:
                    if vMerge.val is None:
                        return ''
                    return vMerge.val
        return None

    @staticmethod
    def is_merge_row(row):
        cells = WordML.get_row_cells(row)
        total = len(cells)
        match = 0
        for col_idx in range(total):
            vMerge = WordML.get_merge_val(row, col_idx)
            if vMerge == "continue":
                match += 1
        if match >= total:
            return True
        return False

    @staticmethod
    def has_vMerge(obj):
        element = obj._element
        vMerge = element.find(".//w:vMerge", namespaces=element.nsmap)
        if vMerge is not None:
            return vMerge.val
    @staticmethod
    def has_vMerge_position(row):
        tr = row._tr
        tc_list = tr.findall(".//w:tc", namespaces=tr.nsmap)
        # print(tc_list)
        for tc_index, tc in enumerate(tc_list):
            vMerge = tc.find(".//w:vMerge", namespaces=tc.nsmap)
            if vMerge is not None:
                return tc_index
    @staticmethod
    def is_merge_restart_cell(cell):
        """
        :param cell: _Cell object
        :return: True is (vMerge has no attribute "restart"| no vMerge)
                False is (vMerge with no attribute)
        """
        tcPr = cell._tc.find(".//w:tcPr", namespaces=cell._tc.nsmap)
        if tcPr is not None:
            vMerge = tcPr.find(".//w:vMerge", namespaces=tcPr.nsmap)
            tcW = tcPr.find(".//w:tcW", namespaces=tcPr.nsmap)
            if vMerge is not None:
                print(vMerge.val)
                if not vMerge.val:
                    return False
                else:
                    if vMerge.val == "restart":
                        return True
        else:
            return True

    @staticmethod
    def insert_paragraph_after(block, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        new_para = Paragraph(new_p, block._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        elm = block._p if WordML.is_paragraph_block(block) else block._tbl
        elm.addnext(new_p)
        return new_para

    @staticmethod
    def duplicate_paragraph_before(block):
        """duplicate a new paragraph after the given paragraph."""
        new_p = copy.deepcopy(block._element)
        new_para = Paragraph(new_p, block._parent)
        elm = block._p if WordML.is_paragraph_block(block) else block._tbl
        elm.addprevious(new_p)
        return new_para

    @staticmethod
    def insert_paragraph_after_table(table, text=None, style=None):
        """Insert a new paragraph after the given table."""
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        new_para = Paragraph(new_p, table._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    @staticmethod
    def insert_heading_after(document, text="heading", level=3, block=None):
        heading = document.add_heading(text, level=level)
        if block is not None:
            elm = block._p if WordML.is_paragraph_block(block) else block._tbl
            elm.addnext(heading._p)
        return heading

    @staticmethod
    def copy_paragraph_after(block, paragraph):
        if isinstance(block, Table):
            WordML.copy_table_after(block, paragraph)
        else:
            pb, p = block._p, paragraph._p
            new_pb = deepcopy(pb)
            p.addnext(new_pb)

    @staticmethod
    def copy_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)

    @staticmethod
    def move_table_after(table, paragraph):
        if isinstance(table, Table):
            tbl, p = table._tbl, paragraph._p
        else:
            tbl, p = table, paragraph._p
        p.addnext(tbl)

    @staticmethod
    def replace_block(old_block, new_block):
        old_elm = old_block._p if WordML.is_paragraph_block(old_block) else old_block._tbl
        new_elm = new_block._p if WordML.is_paragraph_block(new_block) else new_block._tbl
        old_elm.addnext(new_elm)
        WordML.delete_block(old_block)
        return new_block
    @staticmethod
    def insert_block_after(old_block, new_block):
        old_elm = old_block._p if WordML.is_paragraph_block(old_block) else old_block._tbl
        new_elm = new_block._p if WordML.is_paragraph_block(new_block) else new_block._tbl
        old_elm.addnext(new_elm)
        return new_block
    @staticmethod
    def swap_block_list(idx_block1, idx_block2):
        """swap 2 block list"""
        a_elm_list = [ai._p if WordML.is_paragraph_block(ai) else ai._tbl for ai in idx_block1]
        b_elm_list = [bi._p if WordML.is_paragraph_block(bi) else bi._tbl for bi in idx_block2]
        # fix a_list is top list and b_list is down list
        parent = a_elm_list[0].getparent()
        anchor_a = a_elm_list[0].getprevious()
        anchor_b = b_elm_list[0].getprevious()
        for i, elm in enumerate(b_elm_list):
            anchor = parent.index(anchor_a) + 1 if i == 0 else parent.index(b_elm_list[0])
            parent.insert(anchor + i, b_elm_list[i])
        for i, elm in enumerate(a_elm_list):
            anchor = parent.index(anchor_b) + 1 if i == 0 else parent.index(a_elm_list[0])
            parent.insert(anchor + i, a_elm_list[i])
    @staticmethod
    def insert_break_paragraph_after(block):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        new_para = Paragraph(new_p, block._parent)
        run = new_para.add_run()
        r = run._r
        br_elm = parse_xml(
            r'<w:br {} w:type="page"/>'.format(nsdecls('w')))
        r.append(br_elm)
        elm = block._p if WordML.is_paragraph_block(block) else block._tbl
        elm.addnext(new_p)
        return new_para

    @staticmethod
    def get_text_symbol_in_cell(cell):
        text = ''
        for paragraph in cell.paragraphs:
            for r in paragraph.runs:
                for child in r._r.getchildren():
                    if child.tag == (WORD_NAMESPACE + "sym"):
                        char = child.attrib[WORD_NAMESPACE + "char"]
                        if char.startswith("F"):
                            symbol = chr(int(char, 16) - int('F000', 16))
                        else:
                            symbol = chr(int(char, 16))
                        text += symbol
                    if child.tag == (WORD_NAMESPACE + "t"):
                        text += child.text
            text += '\n'
        return text

    @staticmethod
    def is_invalid_line(row):
        count = 0
        cells = WordML.get_row_cells(row)
        total = len(cells)
        for fld_idx in range(len(cells)):
            cell = cells[fld_idx]
            field = cell.text.replace("\n", " ")
            field = field.strip()
            if field == "":
                count += 1
        if count >= 2/4*total:
            return True
        else:
            return False

    @staticmethod
    def get_row_cells(row):
        return [_Cell(tc, tc.getparent()) for tc in ElmHelper.findall(row._tr, 'tc')]
        
    @staticmethod
    def make_paragraph_highlight_in_yellow_and_text_red(paragraph):
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor(255, 0, 0)

    def get_paragraph_numFmt(self, paragraph, numbering_part= None):
        lvl = self.get_paragraph_lvl(paragraph)
        if lvl is not None:
            numFmt = lvl.find(WORD_NAMESPACE + 'numFmt')
            if numFmt is not None:
                numFmt = numFmt.attrib[WORD_NAMESPACE + 'val']
                return numFmt
        else:
            return None

    def get_paragraph_lvl(self, paragraph, numbering_part= None):
        if self.numbering_part is None: return None
        abstractNumId = None
        numId = paragraph._element.find(".//w:pPr/w:numPr/w:numId", namespaces=paragraph._element.nsmap)
        ilvl = paragraph._element.find(".//w:pPr/w:numPr/w:ilvl", namespaces=paragraph._element.nsmap)
        numId_val = None
        ilvl_val = None
        if numId is None:
            regex = r"<w:pStyle w:val=\"(.*)\""
            style_id = ElmHelper.read_elm_attr(paragraph._element, regex, 1)
            result = self.StyleNumIdDict.get(style_id, None)
            if result is not None:
                numId_val, ilvl_val = None, None
                if result["numId"] is not None:
                    numId_val = int(result["numId"])
                    if result["ilvl"] is not None:
                        ilvl_val = int(result["ilvl"])
                    elif result["outlineLvl"] is not None:
                        ilvl_val = int(result["outlineLvl"])
        else:
            numId_val = numId.val
            if ilvl is not None:
                ilvl_val = ilvl.val
            else:
                ilvl_val = None
        ct_numbering = self.numbering_part._element
        if numId_val is not None:
            for num in ct_numbering.num_lst:
                if numId_val == num.numId:  # CT_Num
                    abstractNumId = num.abstractNumId.val
                    break

        if abstractNumId is None or ilvl_val is None:
            return None
        abstractNumlist = self.numbering_part._element.findall(".//w:abstractNum",
                                                               namespaces=self.numbering_part._element.nsmap)
        for abstractNum in abstractNumlist:
            temp_abstractNumId = abstractNum.attrib[WORD_NAMESPACE + 'abstractNumId']
            if temp_abstractNumId is not None:
                temp_abstractNumId = int(abstractNum.attrib[WORD_NAMESPACE + 'abstractNumId'])
                abstractNumId = int(abstractNumId)
                if temp_abstractNumId == abstractNumId:
                    list_lvl = abstractNum.findall(".//w:lvl", namespaces=abstractNum.nsmap)
                    for lvl in list_lvl:
                        temp_ilvl = lvl.attrib[WORD_NAMESPACE + 'ilvl']
                        if temp_ilvl is not None:
                            temp_ilvl = int(temp_ilvl)
                            if temp_ilvl == int(ilvl_val):
                                return lvl
    @staticmethod
    def is_hidden_line(row):
        count = 0
        cells = WordML.get_row_cells(row)
        total = len(cells)
        for fld_idx in range(total):
            cell = cells[fld_idx]
            for para in cell.paragraphs:
                if WordML.is_vanished_paragraph(para):
                    count += 1
                    break
        if count >= 2/4*total:
            return True

    @staticmethod
    def is_vanished_paragraph(paragraph):
        for run in paragraph.runs:
            run = run._r
            rPr = run.find('.//w:rPr', namespaces=run.nsmap)
            if rPr is None:
                continue
            vanish = rPr.find('.//w:vanish', namespaces=rPr.nsmap)
            if vanish is not None:
                return True
        return False

    @staticmethod
    def get_hidden_row_type(row):
        hidden_type = 0
        count = 0
        count_empty = 0
        count_hidden = 0
        cells = WordML.get_row_cells(row)
        total = len(cells)
        for fld_idx in range(total):
            cell = cells[fld_idx]
            for paragraph in cell.paragraphs:
                paragraph_text = WordML.get_full_latest_paragraph_text(paragraph)
                if paragraph_text == "":
                    count_empty += 1
                    continue
                paragraph_runs = WordML.get_full_latest_run_paragraph(paragraph)
                not_disclosed_paragraph = ""
                for run in paragraph_runs:
                    rPr = run.find('.//w:rPr', namespaces=run.nsmap)
                    if rPr is None:
                        continue
                    vanish = rPr.find('.//w:vanish', namespaces=rPr.nsmap)
                    if vanish is not None:
                        not_disclosed_paragraph = not_disclosed_paragraph + run.text
                        count_hidden += 1
                if paragraph_text.strip() == not_disclosed_paragraph.strip():
                    count += 1
        if (count+count_empty) >= (total-1):
            hidden_type = 1
        else:
            if count_hidden > 0:
                hidden_type = 2
        return hidden_type

    @staticmethod
    def get_cell_text_without_hidden_text(cell):
        local_cell = copy.deepcopy(cell)
        for paragraph in local_cell.paragraphs:
            paragraph_text = WordML.get_full_latest_paragraph_text(paragraph)
            paragraph_runs = WordML.get_full_latest_run_paragraph(paragraph)
            for run in paragraph_runs:
                rPr = run.find('.//w:rPr', namespaces=run.nsmap)
                if rPr is None:
                    continue
                vanish = rPr.find('.//w:vanish', namespaces=rPr.nsmap)
                if vanish is not None:
                    run.getparent().remove(run)
                    run = None
        return local_cell.text

    @staticmethod
    def get_cell_hidden_text(cell):
        local_cell = copy.deepcopy(cell)
        hidden_text = ""
        for paragraph in local_cell.paragraphs:
            paragraph_text = WordML.get_full_latest_paragraph_text(paragraph)
            paragraph_runs = WordML.get_full_latest_run_paragraph(paragraph)
            for run in paragraph_runs:
                rPr = run.find('.//w:rPr', namespaces=run.nsmap)
                if rPr is None:
                    continue
                vanish = rPr.find('.//w:vanish', namespaces=rPr.nsmap)
                if vanish is not None:
                    hidden_text = hidden_text + run.text
        return hidden_text

    @staticmethod
    def make_cell_aligment_center(cell):
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def convert2DocX(input_dir, output_dir):
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False # True for debug
        wb = word.Documents.Open(input_dir)
        # toDocxFile = os.path.join(output_dir,os.path.basename(input_dir).replace(".doc",".docx"))
        wb.SaveAs2(output_dir, FileFormat=16) # file format for docx
        wb.Close()
        word.Quit()
        time.sleep(1)

    @staticmethod
    def set_cell_width(cell, width):
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '%s' % width)
        tcW.set(qn('w:type'), 'dxa')
        cell._tc.get_or_add_tcPr().append(tcW)

    @staticmethod
    def set_row_height(row, val):
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '%s' % val)
        trHeight.set(qn('w:hRule'), 'exact')
        row._tr.get_or_add_trPr().append(trHeight)

    def get_all_styles(self):
        collection = []
        for j, block in enumerate(self.iter_block_items(self.document)):
            if WordML.is_paragraph_block(block):
                if block.text.strip() == "":
                    continue
                pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
                if pStyle is not None:
                    text = block.text
                    style = pStyle.val
                    heading = block.style.name
                    no = len(collection)
                    self.styles.append({"block_id":no, 'text': text, 'style': style, 'heading': heading})
        return self.styles

    @staticmethod
    def get_full_inserted_deleted_run_paragraph(p):
        r_lst = []
        r_ins_lst = []
        r_del_lst = []
        temp_r_list = p._p.r_lst
        for child in p._p.getchildren():
            if child.tag == (WORD_NAMESPACE + "ins"):
                ins_runs = child.findall(".//w:r", namespaces=child.nsmap)
                r_ins_lst.extend(ins_runs)
            if child.tag == (WORD_NAMESPACE + "del"):
                del_runs = child.findall(".//w:r", namespaces=child.nsmap)
                r_del_lst.extend(del_runs)
            if child.tag == (WORD_NAMESPACE + "r"):
                if temp_r_list.count(child) >= 1:
                    r_lst.append(child)
        return r_lst, r_ins_lst, r_del_lst

    @staticmethod
    def get_ins_del_elm_in_paragraph(p):
        ins_lst = []
        del_lst = []
        temp_r_list = p._p.r_lst
        for child in p._p.getchildren():
            if child.tag == (WORD_NAMESPACE + "ins"):
                ins_lst.append(child)
            if child.tag == (WORD_NAMESPACE + "del"):
                del_lst.append(child)
        return ins_lst, del_lst

    @staticmethod
    def get_full_inserted_delete_paragraph_text(p, separator=""):
        r_lst, r_ins_lst, r_del_lst = WordML.get_full_inserted_deleted_run_paragraph(p)
        ins_text = ""
        for run in r_ins_lst:
            temp_t_list = run.t_lst
            t_lst = run.findall('.//w:t', namespaces=run.nsmap)  ##ONLY DIRECT CHILD
            for t in t_lst:
                if temp_t_list.count(t) >= 1:
                    ins_text = ins_text + separator + t.text
        del_text = ""
        for run in r_del_lst:
            t_lst = run.findall('.//w:delText', namespaces=run.nsmap)  ##ONLY DIRECT CHILD
            for t in t_lst:
                del_text = del_text + separator + t.text
        return ins_text, del_text

    @staticmethod
    def get_id_date_author(del_ins_elm):
        # w: id = "12"
        # w: author = "HIROSHI INOSE"
        # w: date = "2018-03-30T13:04:00Z"
        r = del_ins_elm
        id, author, date = "", "", ""
        try:
            id = r.attrib[WORD_NAMESPACE + 'id']
            author = r.attrib[WORD_NAMESPACE + 'author']
            date = r.attrib[WORD_NAMESPACE + 'date']
            return id, author, date
        except:
            return id, author, date


    @staticmethod
    def is_having_numberingChange(block):
        """ Check if has attribute "w:numberingChange"
        Params: block (paragraph) :, must be a heading.
        Return: Boolean value.
        """
        if WordML.is_heading(p):
            numberingChange = p._element.find('.//w:numberingChange', namespaces=p._element.nsmap)
            if numberingChange is not None:
                return True
        return False

    @staticmethod
    def get_original_numberingChange(block):
        """
        """
        reg = r"(\w+.)?(%1):(\d+):(\d+):([^%\s]+)?((%2):(\d+):(\d+):([^%\s]+)?)?((%3):(\d+):(\d+):([^%\s]+)?)?((%4):(\d+):(\d+):([^%\s]+)?)?"

        numberingChange = block._element.find('.//w:numberingChange', namespaces=block._element.nsmap)   
        id, author, date , original_numbering = "","","",""
        if numberingChange is not None:
            id = numberingChange.attrib[WORD_NAMESPACE + 'id']
            author = numberingChange.attrib[WORD_NAMESPACE + 'author']
            date = numberingChange.attrib[WORD_NAMESPACE + 'date']
            original = numberingChange.attrib[WORD_NAMESPACE + 'original']
            # if original is not None:
            # handle regex on original
            match =  re.match(reg,original)
            
            if match:
                if match.group(1) is not None:
                    original_numbering += str(match.group(1))
                original_numbering += str(match.group(3))
                if match.group(5) is not None:
                    original_numbering += str(match.group(5))
                if match.group(7):
                    original_numbering += str(match.group(8))
                    if match.group(10) is not None:
                        original_numbering += str(match.group(10))
                    if match.group(12):
                        original_numbering += str(match.group(13))
                        if match.group(15) is not None:
                            original_numbering += str(match.group(15))
                        if match.group(17):
                            original_numbering += str(match.group(18))
                            if match.group(20) is not None:
                                original_numbering += str(match.group(20))
            # else:
            return id, author, date , original_numbering
        else:
            return id, author, date , original_numbering

    @staticmethod
    def create_comparision_file(original_file, revised_file, comparison_file):
        try:
            # note the \\ at the end of the path name to prevent a SyntaxError
            # Create the Application word
            Application = win32com.client.gencache.EnsureDispatch("Word.Application")
            # Compare documents
            Application.CompareDocuments(Application.Documents.Open(original_file),
                                         Application.Documents.Open(revised_file), 
                                         RevisedAuthor=False, 
                                         IgnoreAllComparisonWarnings=True)
            # before saving if you like viewing the document in Print Layout.
            # Otherwise the saved Comparison.docx opens as Web Layout by default (Type = 6).
            Application.ActiveDocument.ActiveWindow.View.Type = 3
            # Save the comparison document as "Comparison.docx"
            # https://stackoverflow.com/questions/9868830/how-to-add-encoding-parameter-to-word-document-saveas-function-in-vbss
            comparison_file = comparison_file.replace('\'', '\\')
            Application.ActiveDocument.SaveAs(FileName=comparison_file,
                                              Encoding=20127)
            # Don't forget to quit your Application
            Application.Quit()
        except:
            pass

    @staticmethod
    def accept_change(filename,output_dir):
        # try:
        Application = win32com.client.gencache.EnsureDispatch("Word.Application")
        Application.Visible = False
        # filename = filename.replace("/","\\\\")
        file = Application.Documents.Open(filename)
        time.sleep(0.1)
        file.Activate()

        # Accept all revisions
        Application.ActiveDocument.Revisions.AcceptAll()
        # Delete all comments
        if Application.ActiveDocument.Comments.Count >= 1:
            Application.ActiveDocument.DeleteAllComments()

        Application.ActiveDocument.TrackRevisions = False  # Maybe not need this (not really but why not)
        # output_dir = output_dir.replace('\'', '\\')
        Application.ActiveDocument.SaveAs(FileName=output_dir)
        file.Close(False)
        Application.Application.Quit()
        time.sleep(0.1)
        # except:
        #     pass

    @staticmethod
    def iter_unique_cells(row):
        """ Inheriting from get_row_cells()
        Generate cells in *row* skipping empty grid cells."""
        try:
            prior_tc = None
            for cell in row.cells:
                this_tc = cell._tc
                if this_tc is prior_tc:
                    continue
                prior_tc = this_tc
                yield cell
        except Exception as e:
            try:
                cells = []
                col_idx = len(row._tr.tc_lst)
                for i in range(col_idx):
                    tc = row._tr.tc_at_grid_col(i)
                    cell = _Cell(tc, tc.getparent())
                    cells.append(cell)
                yield cells
            except ValueError as e:
                print("        INFO: ValueError: get_row_cells since bad format table" + str(e))
                yield cells
            print("Error: get_row_cells 0 " + str(e))
            pass

    @staticmethod
    def cut_paragraph_text(paragraph, cut_text):
        run_list = [child for child in paragraph._p.getchildren() if child.tag == WORD_NAMESPACE + 'r']
        start = 0
        for i in range(len(run_list)):
            if cut_text not in WordML.get_text_run_list(run_list[i:]):
                start = i - 1
                break
        end = len(run_list)
        for i in range(len(run_list), -1, -1):
            if cut_text not in WordML.get_text_run_list(run_list[start:i]):
                end = i+1
                break
        parts = ["part1", "part2", "part3"]
        caption_run_list = run_list[start: end + 1]
        while len(caption_run_list) < len(parts):
            idx = paragraph._p.getchildren().index(caption_run_list[-1])
            r = copy.deepcopy(caption_run_list[-1])
            paragraph._p.insert(idx+1, r)
            caption_run_list.append(r)
        # map run index and run object
        run_dict = dict()
        for i in range(len(parts)):
            run_dict[parts[i]] = caption_run_list[i]
        # for key in parts:
        #     caption_dict[key] = new_caption_dict[key]
        # for key in parts:
        #     if run_dict[key].text != caption_dict[key]:
        #         run_dict[key].text = caption_dict[key]

    @staticmethod
    def get_text_run_list(run_list):
        text = ""
        for run in run_list:
            temp_t_list = run.t_lst
            t_lst = run.findall('.//w:t', namespaces=run.nsmap)
            for t in t_lst:
                if temp_t_list.count(t) >= 1:
                    text = text + t.text
        return text

    @staticmethod
    def has_break_page(block):
        r_list = ElmHelper.findall(block._element, 'r')
        r_count = 0
        for r in r_list:
            br = r.find('.//w:br', namespaces=r.nsmap)
            is_page = ElmHelper.read_elm_attr2(br, "type")
            if is_page == "page":
                r_count += 1
        return r_count > 0

    @staticmethod
    def remove_break_page(block):
        run_list = [child for child in block._p.getchildren() if child.tag == WORD_NAMESPACE + 'r']
        for r in run_list:
            br = r.find('.//w:br', namespaces=r.nsmap)
            if br is not None:
                br.getparent().remove(br)

    @staticmethod
    def get_style_name(block):
        pStyle = block._element.find('.//w:pStyle', namespaces=block._element.nsmap)
        if pStyle is not None:
            return block.style.name
        return None
    
    @staticmethod
    def is_image_block(block):
        if block._element.find('.//w:drawing/wp:inline', namespaces=block._element.nsmap) is not None:
            return True
        if block._element.find('.//w:drawing/wp:anchor', namespaces=block._element.nsmap) is not None:
            return True
        if block._element.find('.//w:object/v:shape', namespaces=block._element.nsmap) is not None:
            return True
        if block._element.find('.//w:object/v:shape', namespaces=block._element.nsmap) is not None:
            return True
        if block._element.findall(".//w:drawing//", namespaces=block._element.nsmap):
            return True
        if block._element.findall(".//v:shape//", namespaces=block._element.nsmap):
            return True
        pStyle = block._element.find(".//v:pStyle", namespaces=block._element.nsmap)
        if pStyle is not None and pStyle.val in ["box"]:
            return True
        elm_list = [("drawing", "w"),
                    ("inline", "wp"),
                    ("anchor", "wp"),
                    ("pic", "pic"),
                    ("AlternateContent", "mc"),
                    ("shape", "v"),
                    ("shapetype", "v"),
                    ("formulas", "v"),
                    ("path", "v"),
                    ("object", "w"),
                    ("textbox", "v"),
                    ("pict", "w"),
                    ("line", "v"),
                    ]
        for (tag, ns) in elm_list:
            if len(ElmHelper.finditer(block._element, tag, ns)) > 0:
                return True
        return False

    @staticmethod
    def is_image_run(r):
        """ 07122020 TrungNH20 updated
        :param r: should be element (r = run._r | r = run._element)
        :return: boolean if contains image xpath
        """
        elm = r.find('.//w:drawing/wp:inline', namespaces=r.nsmap)
        if elm is not None:
            return True
        elm = r.find('.//w:drawing/wp:anchor', namespaces=r.nsmap)
        if elm is not None:
            return True
        elm = r.find('.//w:object/v:shape', namespaces=r.nsmap)
        if elm is not None:
            return True
        elm = r.find('.//w:pict/v:shape', namespaces=r.nsmap)
        if elm is not None:
            return True
        return False

    def get_image_names(self, block):
        image_name = []
        elm1 = block._element.find('.//w:drawing/wp:inline', namespaces=block._element.nsmap)
        elm2 = block._element.find('.//w:drawing/wp:anchor', namespaces=block._element.nsmap)
        if elm1 is not None or elm2 is not None:
            regex = r"a:blip\sr:embed=\"(rId\d+)\""
            img = [match.group(1) for match in re.finditer(regex, str(block._element.xml))]
            img = list(dict.fromkeys(img))
            for r_id in img:
                image_part = self.related_parts[r_id]
                image_name.append(os.path.basename(image_part.partname))
        return image_name

    def get_layout_of_images(self):
        for j, block in enumerate(self.blocks):
            image_names = self.get_image_names(block)
            if len(image_names) > 0:
                self.layout_images.append({"block_id": j, "image": image_names})
        return self.layout_images

    @staticmethod
    def create_file(block_list, output_file, template_file):
        if output_file != "":
            document = Document(template_file)
            for i, block in enumerate(block_list):
                anchor = document.paragraphs[0] if i == 0 else WordML.insert_paragraph_after(anchor, "")
                anchor = WordML.replace_block(anchor, block)
            document.save(output_file)
        return output_file

    @staticmethod
    def create_break_paragraph_after(block):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        new_para = Paragraph(new_p, block._parent)
        run = new_para.add_run()
        r = run._r
        br_elm = parse_xml(
            r'<w:br {} w:type="page"/>'.format(nsdecls('w')))
        r.append(br_elm)
        elm = block._p if WordML.is_paragraph_block(block) else block._tbl
        return new_para

    @staticmethod
    def insert_block_after(old_block, new_block):
        old_elm = old_block._p if WordML.is_paragraph_block(old_block) else old_block._tbl
        new_elm = new_block._p if WordML.is_paragraph_block(new_block) else new_block._tbl
        old_elm.addnext(new_elm)
        return new_block

    @staticmethod
    def insert_block_before(old_block, new_block):
        old_elm = old_block._p if WordML.is_paragraph_block(old_block) else old_block._tbl
        new_elm = new_block._p if WordML.is_paragraph_block(new_block) else new_block._tbl
        old_elm.addprevious(new_elm)
        return new_block

    def has_fldChar(block):
        fldChar = block._element.findall(".//w:fldChar",namespaces=block._element.nsmap)
        if fldChar != []:
            return True
        else:
            return False

    @staticmethod
    def has_instrText(block):
        instrText = block._element.findall(".//w:instrText",namespaces=block._element.nsmap)
        if instrText != []:
            # return instrText
            return True
        else:
            return False

    @staticmethod
    def has_bookmarkStart(block):
        bookmarkStart = block._element.findall(".//w:bookmarkStart", namespaces=block._element.nsmap)
        if bookmarkStart != []:
            # return instrText
            return True
        else:
            return False

    @staticmethod
    def add_fldChar(value='begin'):
        """
        value: begin/separate/end
        """
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), value)
        return fldChar

    @staticmethod
    def add_instrText(field_code=''):
        instrText = OxmlElement('w:instrText')
        instrText.text = field_code
        return instrText

    @staticmethod
    def add_field(r, field_text="", field_code=""):
        """ using in inserters.crossReference
            - r:
            - field_text: text
            - field_code: toggle field_code
        """
        # run = paragraph.add_run()
        # r = run._r
        r.append(WordML.add_fldChar(value='begin'))
        r.append(WordML.add_instrText(field_code=field_code))
        r.append(WordML.add_fldChar(value='separate'))
        r.append(WordML.add_run_text(field_text))
        r.append(WordML.add_fldChar(value='end'))

    @staticmethod
    def add_reference(paragraph, bookmark_name,field_text='',caption=True):
        """
        caption: True => all text are included (eg: Table 32.1)
                False => using for heading, only display number only (eg: 32)
        """
        if caption == True:
            instrText =' REF {} \\h '.format(bookmark_name)
        else:
            instrText = ' REF {} \\r \\h '.format(bookmark_name)

        run = paragraph.add_run()
        r = run._r
        WordML.add_field(r, field_text=field_text,field_code=instrText)

    @staticmethod
    def add_caption(r,field_type="",field_text=['','']):
        """ using in inserters.crossReference
        Add numbering caption to the paragraph
        Params:
            paragraph:
        eg:
            ____ {reference depend on heading 1}.{sequence caption of Table/Figure} ____
            Table {STYLEREF 1 \s}.{SEQ Table \* ARABIC \s 1} => Table 33.1
            Figure {STYLEREF 1 \s}.{SEQ Figure \* ARABIC \s 1} => Figure 33.2
        """
        instrText1 = ' STYLEREF 1 \\s '
        instrText2 = ' SEQ {} \\* ARABIC \\s 1 '.format(field_type)
        # run = paragraph.add_run()
        # r = run._r
        # if field_type:
        #     r.append(WordML.add_run_text(field_type))
        #     r.append(WordML.add_run_text(' '))
        WordML.add_field(r,field_text=field_text[0],field_code=instrText1)
        r.append(WordML.add_run_text('.'))
        WordML.add_field(r,field_text=field_text[1],field_code=instrText2)


    @staticmethod
    def add_bookmark_start(name="test_bookmark"):
        bm = OxmlElement('w:bookmarkStart')
        bm.set(qn('w:id'), '0')
        bm.set(qn('w:name'), name)
        return bm

    @staticmethod
    def add_bookmark_end():
        bmrk = OxmlElement('w:bookmarkEnd')
        bmrk.set(qn('w:id'), '0')
        return bmrk

    @staticmethod
    def add_run_text(text_contents):
        text = docx.oxml.OxmlElement('w:t')
        text.set(qn('xml:space'), 'preserve')
        text.text = text_contents
        return text

    @staticmethod
    def add_bookmark(paragraph, bookmark_name='',field_text=['',''],bookmark_text='' ,caption=True,field_type="Table"):
        """using in inserters.crossReference
        Params:
            - field_text = [x,y] => x.y (eg: ['33','2'] => 33.2)
            - caption:
                True if bookmark is caption
                False if bookmark is heading
            - bookmark_name: the tag w:name in bookmark_start (xml) for reference
            - bookmark_text: the tag w:r/w:t for display text with bookmark inside
        """
        run = paragraph.add_run()
        run.bold = True

        r = run._r  # for reference the following also works: r =  document.element.xpath('//w:r')[-1]
        r.append(WordML.add_bookmark_start(name=bookmark_name))
        if caption:
            WordML.add_caption(r, field_text=field_text, field_type=field_type)
        else:
            r.append(WordML.add_run_text(bookmark_text))
            # run.add_text(bookmark_text)
            # pass
        r.append(WordML.add_bookmark_end())

    # @staticmethod
    # def add_bookmark

    @staticmethod
    def split_run_in_two(paragraph, run, split_index):
        index_in_paragraph = paragraph._p.index(run.element)

        text_before_split = run.text[0:split_index]
        text_after_split = run.text[split_index:]

        run.text = text_before_split
        new_run = paragraph.add_run(text_after_split)
        WordML.copy_format_manual(run, new_run)
        paragraph._p[index_in_paragraph + 1:index_in_paragraph + 1] = [new_run.element]
        return [run, new_run]

    @staticmethod
    def split_run_in_three(paragraph, run, split_start, split_end):
        first_split = WordML.split_run_in_two(paragraph, run, split_end)
        second_split = WordML.split_run_in_two(paragraph, run, split_start)
        return second_split + [first_split[-1]]

    @staticmethod
    def cumulate_runs_into_first_one(paragraph):

        # whole_para_text = ''.join([run.text for run in paragraph.runs])

        for k, run in enumerate(paragraph.runs):
            if k == 0:
                if WordML.is_image_run(run._r):
                    new_run_r = paragraph._element._new_r()
                    run._r.addnext(new_run_r)
                    new_run = Run(new_run_r, run._parent)
                    new_run.text = paragraph.text
                    text = new_run.text
                else:
                    run.text = paragraph.text
                    text = run.text
            else:
                if WordML.is_image_run(run._r):
                    pass
                elif run.text:
                    run.text = ""
                elif run.text == '\t':
                    paragraph.remove(run._r)

        return paragraph

    @staticmethod
    def copy_format_manual(original_run, copy_run):
        copy_run_font = copy_run.font
        original_run_font = original_run.font
        copy_run_font.bold = original_run_font.bold
        copy_run_font.italic = original_run_font.italic
        copy_run_font.underline = original_run_font.underline
        copy_run_font.strike = original_run_font.strike
        copy_run_font.subscript = original_run_font.subscript
        copy_run_font.superscript = original_run_font.superscript
        copy_run_font.size = original_run_font.size
        copy_run_font.highlight_color = original_run_font.highlight_color
        copy_run_font.color.rgb = original_run_font.color.rgb
        copy_run.style.name = original_run.style.name

    @staticmethod
    def move_paragraph_after(move_paragraph, destination_paragraph):
        """
        :param move_paragraph: the paragraph which will be move
        :param destination_paragraph: the
        :return:
        """
        # if isinstance(move_paragraph, Paragraph):
        #     m_p, d_p = move_paragraph, destination_paragraph._p
        # else:
        m_p, d_p = move_paragraph._p, destination_paragraph._p
        d_p.addnext(m_p)

    @staticmethod
    def replace_text(paragraph, search_text, replace_text):
        if search_text in paragraph.text:
            inline = paragraph.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[
                    -1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length],
                                                          str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text

    @staticmethod
    def create_table(tbl_elm, parent):
        return Table(tbl_elm,  parent)

    @staticmethod
    def create_paragraph(p_elm, parent):
        return Paragraph(p_elm,  parent)

    @staticmethod
    def is_numbering(block):
        numpr = block._element.find('.//w:numPr', namespaces=block._element.nsmap)
        if numpr is not None:
            return True
        return False

    @staticmethod
    def is_merge_cell(row):
        value_vMerge = ElmHelper.read_elm(row, './/w:tcPr/w:vMerge', first=True, val=True)
        if value_vMerge:
            return True
        return False

    @staticmethod
    def scan_table_in_cells_textbox(blocks):
        TextBoxTableList = []
        for idx, block in enumerate(blocks):
            tbl_list = []
            tbl_list1 = WordML.get_all_tables_in_table_cell(block, idx)
            tbl_list2 = WordML.get_all_tables_in_textbox(block, idx)
            tbl_list.extend(tbl_list1)
            tbl_list.extend(tbl_list2)
            if len(tbl_list) > 0:
                TextBoxTableList.append([idx, tbl_list])
        return TextBoxTableList

    @staticmethod
    def get_all_tables_in_table_cell(block, i=0):
        tbl_list = []
        for tr in ElmHelper.findall(block._element, "tr"):
            for tc in ElmHelper.findall(tr, "tc"):
                tbl_i_list = ElmHelper.findall(tc, "tbl")
                if len(tbl_i_list) > 0:
                    tbl_list.extend(tbl_i_list)
        table_list = [WordML.create_table(tbl, tbl.getparent()) for tbl in tbl_list]
        return table_list

    @staticmethod
    def get_all_tables_in_textbox(block, i=0):
        tbl_list = []
        for r in ElmHelper.findall(block._element, "r"):
            for txbx in ElmHelper.finditer(r, "txbxContent"):
                tbl_i_list = ElmHelper.findall(txbx, "tbl")
                if len(tbl_i_list) > 0:
                    tbl_list.extend(tbl_i_list)
        table_list = [WordML.create_table(tbl, tbl.getparent()) for tbl in tbl_list]
        return table_list

    def create_styleid_numid_dict(self):
        regex = r"<w:numId w:val=\"(.*)\""
        regex1 = r"<w:outlineLvl w:val=\"(.*)\""
        regex2 = r"<w:ilvl w:val=\"(.*)\""
        StyleNumIdDict = dict()
        for style in self.document.styles:
            styleId = ElmHelper.read_elm_attr2(style._element, "styleId")
            numid = ElmHelper.read_elm_attr(style._element, regex, 1)
            outlineLvl = ElmHelper.read_elm_attr(style._element, regex1, 1)
            ilvl = ElmHelper.read_elm_attr(style._element, regex2, 1)
            result = dict()
            result['numId'] = numid
            result['outlineLvl'] = outlineLvl
            result['ilvl'] = ilvl
            StyleNumIdDict[styleId] = result
        return StyleNumIdDict
