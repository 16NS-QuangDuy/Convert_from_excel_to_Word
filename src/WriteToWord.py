import pandas as pd
from docx import Document
import DataAccess


document = Document()
def write_to_file(index, level, level_down):
    level_ = level[(level["Tên bố"].notna()) & (level["Tên mẹ"].notna()) & (level["Tên bố"] != "") & (
                level["Tên mẹ"] != "")]

    for item, row in level_.iterrows():
        #Lấy thông tin người trong họ
        name = row["Tên"]
        birth = row["Ngày sinh"]
        death = row["Ngày giỗ"]
        place = row["Thường trú"]
        grave = row["Địa chỉ nghĩa trang"]
        child = DataAccess.all_child(level_down, row)
        count_all = DataAccess.all_child(level_down, row).shape[0]
        count = DataAccess.count_male(level_down, row)
        partner = DataAccess.partner(level, row)

        #Ghi dữ liệu vào file Word
        document.add_heading(f"Đời {index}", 1)
        document.add_paragraph(row["Tên"], style='List Number')
        if not (pd.isna(birth) & pd.isna(death)):
            document.add_paragraph(f"{name}({birth} - {death})")
        else:
            document.add_paragraph(f"{name}")
        if not pd.isna(place):
            document.add_paragraph(f", sinh sống tại {place}, ")
        if not pd.isna(grave):
            document.add_paragraph(f", địa chỉ nghĩa trang tại {grave}.")