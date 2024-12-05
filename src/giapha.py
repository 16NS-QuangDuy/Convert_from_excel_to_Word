import pandas as pd
from docx.shared import Pt

import DataAccess
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from services.wordML import WordML
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_ROW_HEIGHT_RULE

document = Document()
file_path = "Test_Data/Input/TC023/data_test.csv"
df = pd.read_csv(file_path)
df = df.fillna("")


# Hàm để thêm viền cho bảng
def add_table_border(table):
    tbl = table._tbl  # Lấy đối tượng XML của bảng

    # Kiểm tra nếu tblBorders không tồn tại, tạo mới phần tử tblBorders
    tblBorders = tbl.xpath(".//w:tblBorders")
    if not tblBorders:
        tblBorders = OxmlElement("w:tblBorders")
        tbl.tblPr.append(tblBorders)

    # Cấu hình viền cho bảng (đặt kiểu viền, kích thước và khoảng cách)
    for border_name in ['top', 'left', 'right', 'bottom', 'insideV']:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Đặt kích thước viền
        border.set(qn('w:space'), '0')  # Khoảng cách viền

        tblBorders.append(border)  # Thêm viền vào phần tử tblBorders

# Lặp qua danh sách dữ liệu để điền vào bảng
def write_to_table(index, row_user, row_partner):
    if row_partner is not None and not row_partner.empty:
        table = document.add_table(rows=1, cols=2)
        for i, (_, row) in enumerate(row_partner.iterrows()):
            if i != row_partner.shape[0] - 1:
                table.add_row()
            # Điền dữ liệu vào cột 1 (Thông tin vợ/chồng)
            content = f"{row['Tên'].upper()}"
            # Chỉ thêm dòng nếu giá trị không rỗng
            if row["Ngày sinh"]:
                content += f"\nSinh: {row['Ngày sinh']}"
            if row["Ngày giỗ"]:
                content += f"\nMất: {row['Ngày giỗ']} (ÂL)"
            if row["Nghĩa trang"]:
                content += f"\nMộ: {row['Nghĩa trang']}"
            table.cell(0 + i, 1).text = content
            WordML.make_cell_aligment_center(table.cell(0 + i, 1))
            for paragraph in table.cell(0 + i, 1).paragraphs:
                paragraph.paragraph_format.space_before = Pt(10)
            for paragraphIndex, paragraph in enumerate(table.cell(0 + i, 1).paragraphs):
                # Lấy nội dung của đoạn văn
                text = paragraph.text
                if text:
                    # Tách các dòng trong đoạn văn
                    lines = text.split('\n')
                    if lines:
                        # Xác định dòng đầu tiên
                        first_line = lines[0]

                        # Làm rỗng đoạn văn để tùy chỉnh từng dòng
                        paragraph.clear()

                        # Thêm dòng đầu tiên với định dạng bôi đậm
                        run = paragraph.add_run(first_line)
                        run.bold = True

                        # Thêm các dòng còn lại (không bôi đậm)
                        for line in lines[1:]:
                            paragraph.add_run(f'\n{line}')
            WordML.set_cell_font_size(table.cell(0 + i, 1), 'Times New Roman', 14 )
    else:
        table = document.add_table(rows=1, cols=1)

    # Điền dữ liệu vào cột 1 (Thông tin con)
    content = f"{index}{row_user['Tên'].upper()}"
    # Chỉ thêm dòng nếu giá trị không rỗng
    if row_user["Ngày sinh"]:
        content += f"\nSinh: {row_user['Ngày sinh']}"
    if row_user["Ngày giỗ"]:
        content += f"\nMất: {row_user['Ngày giỗ']} (ÂL)"
    if row_user["Nghĩa trang"]:
        content += f"\nMộ: {row_user['Nghĩa trang']}"
    table.cell(0, 0).text = content
    WordML.make_cell_aligment_center(table.cell(0, 0))
    for paragraph in table.cell(0, 0).paragraphs:
        paragraph.paragraph_format.space_before = Pt(10)
    for paragraphIndex, paragraph in enumerate(table.cell(0, 0).paragraphs):
        # Lấy nội dung của đoạn văn
        text = paragraph.text
        if text:
            # Tách các dòng trong đoạn văn
            lines = text.split('\n')
            if lines:
                # Xác định dòng đầu tiên
                first_line = lines[0]

                # Làm rỗng đoạn văn để tùy chỉnh từng dòng
                paragraph.clear()

                # Thêm dòng đầu tiên với định dạng bôi đậm
                run = paragraph.add_run(first_line)
                run.bold = True

                # Thêm các dòng còn lại (không bôi đậm)
                for line in lines[1:]:
                    paragraph.add_run(f'\n{line}')
    WordML.set_cell_font_size(table.cell(0, 0), 'Times New Roman', 14)

    # Thay đổi kích thước của các ô
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(2.31)  # Đặt chiều rộng của ô thành 1.5 inch

    # Thay đổi chiều cao của hàng
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Tự động điều chỉnh chiều cao theo nội dung
        row.height = None  # Bỏ giá trị chiều cao cố định nếu có

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm viền cho bảng mới tạo
    add_table_border(table)

def write_to_file(life, level):
    d = document.add_heading(f"ĐỜI THỨ {life}", 1)
    d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    level_ = level[(level["Tên bố"] != "") | (level["Tên mẹ"] != "")]
    grouped = level_.groupby(["Tên bố", "Tên mẹ"])
    # Tạo danh sách các DataFrame con cho mỗi nhóm
    dataframes = [group for _, group in grouped]

    if life != 1:
        for i, group_df in enumerate(dataframes):
            group_df_sort = group_df.sort_values(by = "STT")
            content = ""
            if group_df_sort.iloc[0]["Tên bố"]:
                content = f"CON ÔNG {group_df_sort.iloc[0]["Tên bố"].upper()}"
            if group_df_sort.iloc[0]["Tên mẹ"]:
                content += f" VÀ BÀ {group_df_sort.iloc[0]["Tên mẹ"].upper()}"
            paragraph = document.add_paragraph()
            run1 = paragraph.add_run(content)
            run1.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for index, (_, row_person) in enumerate(group_df_sort.iterrows()):
                partner = DataAccess.partner(level, row_person)
                write_to_table(f"{index + 1}. ", row_person, partner)
                paragraphs = row_person["Đầy đủ"].split('\n')
                for para in paragraphs:
                    d = document.add_paragraph(para)
                    d.paragraph_format.space_before = Pt(10)
                    d.paragraph_format.first_line_indent = Inches(0.5)
    else:
        row_person = level[level["Giới tính"] == "Nam"].iloc[0]
        partner = DataAccess.partner(level, row_person)
        write_to_table("", row_person, partner)
        paragraphs = row_person["Đầy đủ"].split('\n')
        for para in paragraphs:
            d = document.add_paragraph(para)
            d.paragraph_format.space_before = Pt(10)
            d.paragraph_format.first_line_indent = Inches(0.5)

    # Đổi font chữ cho toàn bộ nội dung
    for paragraph in document.paragraphs:  # Duyệt qua từng đoạn văn
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:  # Duyệt qua từng 'run' trong đoạn
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

for life in df["Đời"].unique():
    level = df[(df["Đời"] == life)]
    write_to_file(life, level)
document.save('output/giapha.docx')
