from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Mở file Word
doc = Document("giapha2.docx")

# Dữ liệu mẫu để điền vào bảng (thêm người con thứ 2)
data = [
    {
        "con": {
            "tên": "Nguyễn Văn A",
            "năm_sinh_mất": "1910 - 1985",
            "ngày_giỗ": "01/01",
            "tên_mộ": "Mộ Ông A"
        },
        "vợ_chồng": {
            "tên": "Trần Thị B",
            "năm_sinh_mất": "1915 - 1990",
            "ngày_giỗ": "02/02",
            "tên_mộ": "Mộ Bà B"
        }
    },
    {
        "con": {
            "tên": "Nguyễn Văn C",
            "năm_sinh_mất": "1940 - 2020",
            "ngày_giỗ": "03/03",
            "tên_mộ": "Mộ Ông C"
        },
        "vợ_chồng": {
            "tên": "Lê Thị D",
            "năm_sinh_mất": "1945 - 2021",
            "ngày_giỗ": "04/04",
            "tên_mộ": "Mộ Bà D"
        }
    }
]

# Kiểm tra số bảng trong tài liệu
tables = doc.tables

# Nếu bảng chưa có, thêm bảng đầu tiên
if not tables:
    table = doc.add_table(rows=1, cols=2)
else:
    # Nếu đã có bảng, lấy bảng đầu tiên
    table = tables[0]

# Hàm để thêm viền cho bảng
def add_table_border(table):
    tbl = table._tbl  # Lấy đối tượng XML của bảng
    
    # Kiểm tra nếu tblBorders không tồn tại, tạo mới phần tử tblBorders
    tblBorders = tbl.xpath(".//w:tblBorders")
    if not tblBorders:
        tblBorders = OxmlElement("w:tblBorders")
        tbl.tblPr.append(tblBorders)

    # Cấu hình viền cho bảng (đặt kiểu viền, kích thước và khoảng cách)
    for border_name in ['top', 'left', 'right', 'bottom', 'insideH', 'insideV']:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Đặt kích thước viền
        border.set(qn('w:space'), '0')  # Khoảng cách viền

        tblBorders.append(border)  # Thêm viền vào phần tử tblBorders

# Lặp qua danh sách dữ liệu để điền vào bảng
for index, record in enumerate(data):
    # Nếu là người con thứ 2 trở đi, tạo bảng mới
    if index > 0:
        table = doc.add_table(rows=1, cols=2)  # Tạo bảng mới cho người con thứ 2 và các con tiếp theo

    # Thêm 4 hàng mới vào bảng
    for _ in range(4):
        table.add_row()

    # Tính vị trí hàng đầu tiên trong nhóm 4 hàng mới
    start_row = len(table.rows) - 4

    # Điền dữ liệu vào cột 1 (Thông tin con)
    table.cell(start_row, 0).text = record["con"]["tên"]
    table.cell(start_row + 1, 0).text = record["con"]["năm_sinh_mất"]
    table.cell(start_row + 2, 0).text = record["con"]["ngày_giỗ"]
    table.cell(start_row + 3, 0).text = record["con"]["tên_mộ"]

    # Điền dữ liệu vào cột 2 (Thông tin vợ/chồng)
    table.cell(start_row, 1).text = record["vợ_chồng"]["tên"]
    table.cell(start_row + 1, 1).text = record["vợ_chồng"]["năm_sinh_mất"]
    table.cell(start_row + 2, 1).text = record["vợ_chồng"]["ngày_giỗ"]
    table.cell(start_row + 3, 1).text = record["vợ_chồng"]["tên_mộ"]

    # Thêm viền cho bảng mới tạo
    add_table_border(table)

# Lưu tài liệu mới
doc.save("giapha2_updated_with_borders.docx")
print("Dữ liệu đã được điền thành công và lưu vào 'giapha2_updated_with_borders.docx'")
